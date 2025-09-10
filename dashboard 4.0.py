# app_b4_1.py — Streamlit Funding Dashboard (Schuman-branded)
from __future__ import annotations
import io, re, base64
from datetime import datetime
from typing import List, Dict, Optional
from pathlib import Path

import pandas as pd
import streamlit as st
import altair as alt

# Optional DOCX for shortlist export
try:
    from docx import Document
    from docx.shared import Pt, Inches   # Inches for image sizing
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# For exporting a PNG Gantt into the DOCX
import matplotlib
matplotlib.use("Agg")  # headless backend
import matplotlib.pyplot as plt
import matplotlib.dates as mdates

# ------------------------------------------------------------
# Brand assets locations
ASSETS_DIR = Path("assets")
FONTS_DIR  = ASSETS_DIR / "fonts"
LOGO_BLUE  = ASSETS_DIR / "logo-schuman_blue.png"
LOGO_GREY  = ASSETS_DIR / "logo-schuman_grey.png"
LOGO_WHITE = ASSETS_DIR / "logo-schuman_white.png"
LOGO_SMALL = ASSETS_DIR / "logo-schuman-blue-white-bg-SMALL.png"

FONT_FILES = [
    ("SA Brand", "normal", 300, ASSETS_DIR / "Frm-Light.otf"),
    ("SA Brand", "normal", 400, ASSETS_DIR / "Frm-Regular.otf"),
    ("SA Brand", "normal", 500, ASSETS_DIR / "Frm-Medium.otf"),
    ("SA Brand", "normal", 700, ASSETS_DIR / "Frm-Bold.otf"),
    ("SA Brand", "italic", 700, ASSETS_DIR / "Frm-Bold-Italic.otf"),
    ("SA Brand", "normal", 900, ASSETS_DIR / "Frm-Black.otf"),
]

# ---------- small utils ----------

# Make dates robust for Gantt (handles missing/invalid ranges)
def _prepare_dates_for_chart(df: pd.DataFrame, default_window_days: int = 60) -> pd.DataFrame:
    g = df.copy()

    # End date: prefer final deadline; fall back to closing_date_any if present
    end = g.get("deadline")
    if "closing_date_any" in g.columns:
        end = end.fillna(g["closing_date_any"]) if end is not None else g["closing_date_any"]
    g["_chart_end"] = end

    # Start date: prefer opening_date; otherwise synthesize as end - default_window_days
    start = g.get("opening_date")
    g["_chart_start"] = start
    have_end_only = g["_chart_start"].isna() & g["_chart_end"].notna()
    g.loc[have_end_only, "_chart_start"] = g.loc[have_end_only, "_chart_end"] - pd.Timedelta(days=default_window_days)

    # Drop rows with no usable range
    g = g[g["_chart_start"].notna() & g["_chart_end"].notna()].copy()

    # Fix inverted ranges (start after end) by clipping start to a week before end
    inverted = g["_chart_start"] > g["_chart_end"]
    g.loc[inverted, "_chart_start"] = g.loc[inverted, "_chart_end"] - pd.Timedelta(days=7)

    # The chart builder expects columns named opening_date/deadline
    g["opening_date"] = g["_chart_start"]
    g["deadline"]     = g["_chart_end"]
    return g

@st.cache_data(show_spinner=False)
def _file_to_base64(p: Path) -> str | None:
    try:
        return base64.b64encode(p.read_bytes()).decode("utf-8")
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def _read_theme_css(path: Path) -> str:
    # include mtime in cache key to invalidate on edit
    _ = path.stat().st_mtime if path.exists() else 0
    return path.read_text(encoding="utf-8") if path.exists() else ""

def inject_brand_css():
    """Injects @font-face (dynamic from OTFs) + static theme from assets/theme.css."""
    # 1) Build @font-face blocks for any available OTFs (base64 inline)
    font_faces = []
    for fam, style, weight, path in FONT_FILES:
        if path.exists():
            b64 = _file_to_base64(path)
            if b64:
                font_faces.append(f"""
                @font-face {{
                  font-family: '{fam}';
                  src: url(data:font/otf;base64,{b64}) format('opentype');
                  font-style: {style};
                  font-weight: {weight};
                  font-display: swap;
                }}
                """)
    font_css = "\n".join(font_faces)

    # 2) Load static theme css
    theme_css = _read_theme_css(ASSETS_DIR / "theme.css")

    # 3) Inject
    st.markdown(f"<style>\n{font_css}\n{theme_css}\n</style>", unsafe_allow_html=True)

def brand_header():
    """Top hero header with gradient and logo."""
    logo_src = None
    for p in (LOGO_WHITE, LOGO_GREY, LOGO_BLUE):
        if p.exists():
            logo_src = f"data:image/png;base64,{_file_to_base64(p)}"
            break

    st.markdown(f"""
    <div style="
      border-radius: 16px;
      background: var(--sa-primary);
      padding: 24px 20px;
      color: white;
      text-align: center;
      box-shadow: var(--sa-shadow);">
      {'<img src="'+logo_src+'" alt="Schuman Associates" style="height:60px; margin-bottom:12px;" />' if logo_src else ''}
      <div style="font-size:20px; font-weight:700; margin-bottom:4px;">Schuman Associates · Funding Dashboard</div>
      <div style="font-size:14px; opacity:.9;">Your European partners in a global market since 1989</div>
    </div>
    """, unsafe_allow_html=True)

# ------------------------------------------------------------
# Column mapping / helpers / canonicalisation
COLUMN_MAP = {
    "Programme": "programme", "Code": "code", "Title": "title",
    "Opening Date": "opening_date", "Opening date": "opening_date",
    "Deadline": "deadline", "First Stage Deadline": "first_deadline",
    "Second Stage Deadline": "second_deadline", "Second Stage deadline": "second_deadline",
    "Two-Stage": "two_stage", "Cluster": "cluster", "Destination": "destination",
    "Destination / Strand": "destination", "Destination/Strand": "destination", "Strand": "destination",
    "Budget Per Project": "budget_per_project_eur", "Budget per project": "budget_per_project_eur",
    "Total Budget": "total_budget_eur", "Number of Projects": "num_projects",
    "Type of Action": "type_of_action", "TRL": "trl", "Call Name": "call_name",
    "Expected Outcome": "expected_outcome", "Scope": "scope", "Description": "full_text",
    "Source Filename": "source_filename", "Version Label": "version_label", "Parsed On (UTC)": "parsed_on_utc",
    # Erasmus-specific:
    "Managing Authority": "managing_authority", "Key Action": "key_action",
}
DISPLAY_COLS = [
    "programme","code","title","opening_date","deadline",
    "type_of_action","budget_per_project_eur",
    "cluster","destination","trl",
    "managing_authority","key_action",
    "first_deadline","second_deadline","two_stage",
    "call_name","version_label","source_filename","parsed_on_utc",
]
SEARCHABLE_COLUMNS = (
    "code","title","call_name","expected_outcome","scope","full_text",
    "cluster","destination","type_of_action","trl","managing_authority","key_action"
)

def nl_to_br(s: str) -> str:
    return "" if not s else s.replace("\n", "<br>")

def clean_footer(text: str) -> str:
    if not text:
        return ""
    pat = re.compile(r"Horizon\s*Europe\s*[-–]?\s*Work Programme.*?Page\s+\d+\s+of\s+\d+", re.IGNORECASE | re.DOTALL)
    cleaned = pat.sub("", text)
    return re.sub(r"\s+", " ", cleaned).strip()

def normalize_bullets(text: str) -> str:
    if not isinstance(text, str) or text == "":
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?m)^[ \t]*[▪◦●•]\s*", "- ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"(?<!\n)([ \t]+[-*]\s+)", r"\n- ", text)
    text = re.sub(r"(?<!\n)([ \t]+)(\d+\.\s+)", r"\n\2", text)
    return text.strip()

def highlight_text(text: str, keywords: list[str], colours=None) -> str:
    if not text:
        return ""
    kws = [str(k).strip() for k in keywords if k and str(k).strip()]
    if not kws:
        return text
    if colours is None:
        colours = ["#ffff00", "#a0e7e5", "#ffb3b3"]
    out = str(text)
    for i, kw in enumerate(kws):
        colour = colours[i % len(colours)]
        out = re.sub(re.escape(kw), lambda m: f"<span style='background-color:{colour}; font-weight:bold;'>{m.group(0)}</span>", out, flags=re.IGNORECASE)
    return out

def safe_date_series(s: pd.Series) -> pd.Series:
    out = pd.to_datetime(s, errors="coerce", dayfirst=True)
    if out.notna().sum() == 0:
        out = pd.to_datetime(s, errors="coerce", dayfirst=False)
    return out

def canonicalise(df: pd.DataFrame, programme_name: str) -> pd.DataFrame:
    df.columns = [c.strip() for c in df.columns]
    for src, dst in COLUMN_MAP.items():
        if src in df.columns and dst not in df.columns:
            df = df.rename(columns={src: dst})
    df = df.rename(columns={c: c.strip().lower() for c in df.columns})

    # programme comes from the sheet we're loading
    df["programme"] = programme_name

    for c in ("budget_per_project_eur","total_budget_eur","trl","num_projects"):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    for c in ("opening_date","deadline","first_deadline","second_deadline"):
        if c in df.columns:
            df[c] = safe_date_series(df[c])

    if "two_stage" in df.columns:
        df["two_stage"] = (
            df["two_stage"].astype(str).str.strip().str.lower()
            .map({"true": True, "false": False, "yes": True, "no": False, "1": True, "0": False})
            .fillna(False)
        )
    else:
        df["two_stage"] = False

    present = [c for c in SEARCHABLE_COLUMNS if c in df.columns]
    df["_search_all"]   = df[present].astype(str).agg(" ".join, axis=1).str.lower() if present else ""
    title_cols = [c for c in ["code","title"] if c in df.columns]
    df["_search_title"] = df[title_cols].astype(str).agg(" ".join, axis=1).str.lower() if title_cols else ""

    close_cols = [c for c in ["deadline","first_deadline","second_deadline"] if c in df.columns]
    if close_cols:
        df["closing_date_any"] = pd.to_datetime(df[close_cols].stack(), errors="coerce").groupby(level=0).min()
    else:
        df["closing_date_any"] = pd.NaT
    df["opening_year"]  = df["opening_date"].dt.year
    df["deadline_year"] = df["deadline"].dt.year

    return df

# --------- Caching ----------
@st.cache_data(show_spinner=False)
def get_sheet_names(file_bytes: bytes) -> List[str]:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    return xls.sheet_names

@st.cache_data(show_spinner=False)
def load_programme(file_bytes: bytes, sheet_name: str, programme_name: str, _ver:int=1) -> pd.DataFrame:
    _ = hash(file_bytes)  # cache key includes content
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    raw = pd.read_excel(xls, sheet_name=sheet_name)
    df = canonicalise(raw, programme_name)
    return df.copy(deep=True)

# --------- Chart prep ----------
def wrap_label(text: str, width=60, max_lines=3) -> str:
    s = str(text or "")
    chunks = [s[i:i+width] for i in range(0, len(s), width)]
    return "\n".join(chunks[:max_lines]) if chunks else "—"

def build_month_bands(min_x: pd.Timestamp, max_x: pd.Timestamp) -> pd.DataFrame:
    start = pd.Timestamp(min_x).to_period("M").start_time
    end   = (pd.Timestamp(max_x).to_period("M") + 1).start_time
    months = pd.date_range(start, end, freq="MS")
    return pd.DataFrame({"start": months[:-1], "end": months[1:], "band": [i % 2 for i in range(len(months)-1)]})

def build_singlebar_rows(df: pd.DataFrame) -> pd.DataFrame:
    g = df.copy()
    if "code" in g.columns and g["code"].notna().any():
        base = g["code"].fillna("").astype(str)
    elif "title" in g.columns and g["title"].notna().any():
        base = g["title"].fillna("").astype(str)
    else:
        base = pd.Series([f"row-{i}" for i in range(len(g))], index=g.index)
    if base.duplicated(keep=False).any():
        base = base + g.groupby(base).cumcount().astype(str).radd("#")

    g["y_label"] = base.map(lambda s: wrap_label(s, width=100, max_lines=5))
    g["title_inbar"] = g.get("title","").astype(str).map(lambda s: wrap_label(s, width=100, max_lines=3))
    g = g[pd.notna(g["opening_date"]) & pd.notna(g["deadline"]) & (g["opening_date"] <= g["deadline"])].copy()
    if g.empty:
        return g
    g["bar_days"] = (g["deadline"] - g["opening_date"]).dt.days
    g["mid"] = g["opening_date"] + (g["deadline"] - g["opening_date"])/2
    return g.sort_values(["deadline","opening_date"])

def sa_altair_theme():
    # Align charts with UI fonts & colors
    return {
        "config": {
            "font": "SA Brand",
            "axis":   {"labelFont": "SA Brand", "labelFontWeight": 500, "titleFont": "SA Brand", "labelColor":"#0F172A"},
            "legend": {"labelFont": "SA Brand", "titleFont": "SA Brand"},
            "header": {"labelFont": "SA Brand", "titleFont": "SA Brand"},
            "title":  {"font": "SA Brand", "fontSize": 16, "fontWeight": 700, "color":"#0F172A"},
            "range": {
                "category": ["#1E4F86","#66C2A5","#FC8D62","#8DA0CB","#E78AC3","#A6D854","#FFD92F","#E5C494","#B3B3B3"]
            },
            "view": {"stroke": "transparent"}
        }
    }

alt.themes.register("sa_theme", sa_altair_theme)
alt.themes.enable("sa_theme")

def gantt_singlebar_chart(g: pd.DataFrame, color_field: str = "type_of_action", title: str = ""):
    if g is None or g.empty:
        return None

    min_x = min(g["opening_date"].min(), g["deadline"].min())
    max_x = max(g["opening_date"].max(), g["deadline"].max())
    bands_df = build_month_bands(min_x, max_x)

    month_shade = alt.Chart(bands_df).mark_rect(tooltip=False).encode(
        x="start:T", x2="end:T",
        opacity=alt.Opacity("band:Q", scale=alt.Scale(domain=[0,1], range=[0.0,0.05]), legend=None),
        color=alt.value("#1E4F86")
    )
    months = pd.date_range(pd.Timestamp(min_x).to_period("M").start_time,
                           pd.Timestamp(max_x).to_period("M").end_time, freq="MS")
    month_grid = alt.Chart(pd.DataFrame({"t": months})).mark_rule(stroke="#1E4F86", strokeWidth=0.3).encode(x="t:T")
    month_labels_df = pd.DataFrame({
        "month": months[:-1], "next_month": months[1:],
        "label": [m.strftime("%b %Y") for m in months[:-1]]
    })
    month_labels_df["mid"] = month_labels_df["month"] + ((month_labels_df["next_month"] - month_labels_df["month"]) / 2)
    month_labels = alt.Chart(month_labels_df).mark_text(
        align="center", baseline="top", dy=0, fontSize=12, fontWeight="bold"
    ).encode(x="mid:T", text="label:N", y=alt.value(0))

    today_ts = pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None)
    today_df = pd.DataFrame({"t":[today_ts]})
    today_rule = alt.Chart(today_df).mark_rule(color="#1E4F86", strokeDash=[2,1], strokeWidth=2).encode(
        x="t:T", tooltip=[alt.Tooltip("t:T", title="Today", format="%d %b %Y")]
    )
    today_label = alt.Chart(today_df).mark_text(
        align="left", baseline="top", dx=4, dy=18, fontSize=11, fontWeight="bold", color="#1E4F86"
    ).encode(x="t:T", y=alt.value(0), text=alt.Text("t:T", format='Today: "%d %b %Y"'))

    y_order = g["y_label"].drop_duplicates().tolist()
    row_h = 46
    bar_size = int(row_h * 0.38)
    domain_min, domain_max = g["opening_date"].min(), g["deadline"].max()

    base = alt.Chart(g).encode(
        y=alt.Y("y_label:N", sort=y_order,
                axis=alt.Axis(title=None, labelLimit=200, labelFontSize=11, labelAlign="right", labelPadding=50, domain=True),
                scale=alt.Scale(domain=y_order, paddingInner=0.3, paddingOuter=0.8))
    )

    bars = base.mark_bar(cornerRadius=10, size=bar_size).encode(
        x=alt.X("opening_date:T",
                axis=alt.Axis(title=None, format="%b %Y", tickCount="month", orient="top",
                              labelFontSize=11, labelPadding=50, labelOverlap="greedy", tickSize=6),
                scale=alt.Scale(domain=[domain_min, domain_max])),
        x2="deadline:T",
        color=alt.Color(color_field + ":N",
                        legend=alt.Legend(title=color_field.replace("_"," ").title(),
                                          orient="top", direction="horizontal", offset=100),
                        scale=alt.Scale(scheme="set2")),
        tooltip=[
            alt.Tooltip("title:N", title="Title"),
            alt.Tooltip("opening_date:T", title="Opening", format="%d %b %Y"),
            alt.Tooltip("deadline:T", title="Deadline", format="%d %b %Y"),
            alt.Tooltip(color_field + ":N", title=color_field.replace("_"," ").title()),
        ]
    )

    start_labels = base.mark_text(align="right", dx=-4, dy=5, fontSize=10, color="#111")\
        .encode(x="opening_date:T", text=alt.Text("opening_date:T", format="%d %b %Y"))
    end_labels   = base.mark_text(align="left",  dx=4, dy=5, fontSize=10, color="#111")\
        .encode(x="deadline:T",      text=alt.Text("deadline:T",      format="%d %b %Y"))
    
    inbar = base.mark_text(align="left",
                           baseline="bottom",
                           dx=2, 
                           dy=-(int(bar_size/2)+4),
                           color="black",
                           font="SA Brand",
                           fontSize=12,
                           fontWeight=700)\
        .encode(x=alt.X("opening_date:T",
                        scale=alt.Scale(domain=[domain_min, domain_max]),
                        axis=None),
                text="title_inbar:N",
                opacity=alt.condition(alt.datum.bar_days >= 10, alt.value(1), alt.value(0)))

    chart = (month_shade + month_grid + bars + start_labels + end_labels + inbar + month_labels + today_rule + today_label)\
        .properties(height=max(800, len(y_order)*row_h), width='container',
                    padding={"top":50,"bottom":30,"left":10,"right":10})\
        .configure_axis(grid=False, domain=True, domainWidth=1)\
        .configure_view(continuousHeight=500, continuousWidth=500, strokeWidth=0, clip=False)\
        .interactive(bind_x=True)

    return chart if not title else chart.properties(title=title)

# --------- DOCX + Gantt PNG ----------
def _shortlist_gantt_png(df: pd.DataFrame, color_by: str = "type_of_action") -> Optional[bytes]:
    """
    Render a compact horizontal Gantt PNG for the shortlist to embed into DOCX.
    Uses matplotlib (no altair_saver/node).
    """
    if df is None or df.empty:
        return None
    g = _prepare_dates_for_chart(df)
    g = g[pd.notna(g["opening_date"]) & pd.notna(g["deadline"]) & (g["opening_date"] <= g["deadline"])].copy()
    if g.empty:
        return None

    # Labels (prefer code then title), ensure uniqueness
    base = g["code"].fillna("").astype(str)
    fallback = g["title"].fillna("").astype(str)
    labels = base.where(base.ne(""), fallback)
    labels = labels + g.groupby(labels).cumcount().replace(0, "").astype(str).radd("#").replace("#0", "")

    g["_y"] = labels
    g = g.sort_values(["deadline","opening_date"]).reset_index(drop=True)

    # Colors by 'color_by' if present
    categories = g.get(color_by) if color_by in g.columns else None
    if categories is not None:
        cats = categories.fillna("—").astype(str)
        uniq = pd.unique(cats)
        cmap = plt.get_cmap("tab20")
        color_map = {c: cmap(i % 20) for i, c in enumerate(uniq)}
        colors = cats.map(color_map)
    else:
        colors = None

    n = len(g)
    h = max(3.0, min(0.4 * n + 1.0, 10.0))  # ~0.4in/row
    fig, ax = plt.subplots(figsize=(10, h), dpi=160)

    start_nums = mdates.date2num(pd.to_datetime(g["opening_date"]).dt.to_pydatetime())
    dur = (pd.to_datetime(g["deadline"]) - pd.to_datetime(g["opening_date"])).dt.days.clip(lower=1)

    y_pos = range(n)
    ax.barh(list(y_pos), dur, left=start_nums, height=0.35,
            align="center", color=(colors if colors is not None else None), edgecolor="none")

    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(g["_y"])
    ax.invert_yaxis()

    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))

    today = mdates.date2num(pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None).to_pydatetime())
    ax.axvline(today, linestyle="--", linewidth=1.5, color="#1E4F86")

    ax.grid(axis="x", linestyle=":", linewidth=0.5, alpha=0.6)
    ax.set_xlabel("")
    ax.set_ylabel("")
    plt.tight_layout()

    bio = io.BytesIO()
    plt.savefig(bio, format="png", bbox_inches="tight")
    plt.close(fig)
    bio.seek(0)
    return bio.getvalue()

def generate_docx_report(calls_df: pd.DataFrame, notes_by_code: Dict[str,str], title="Funding Report",
                         shortlist_gantt_png: Optional[bytes] = None) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    h = doc.add_heading(title, level=0); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"Generated on {datetime.utcnow():%d %b %Y, %H:%M UTC}"); p.runs[0].font.size = Pt(9)

    # Include shortlist Gantt image if provided
    if shortlist_gantt_png:
        doc.add_heading("Shortlist Gantt", level=1)
        doc.add_picture(io.BytesIO(shortlist_gantt_png), width=Inches(6.5))

    table = doc.add_table(rows=1, cols=5); hdr = table.rows[0].cells
    for i, t in enumerate(["Programme","Code","Title","Opening","Deadline"]): hdr[i].text = t
    for _, r in calls_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r.get("programme","")); row[1].text = str(r.get("code","")); row[2].text = str(r.get("title",""))
        op, dl = r.get("opening_date"), r.get("deadline")
        row[3].text = op.strftime("%d %b %Y") if pd.notna(op) else "-"
        row[4].text = dl.strftime("%d %b %Y") if pd.notna(dl) else "-"

    for _, r in calls_df.iterrows():
        doc.add_page_break()
        doc.add_heading(f"{r.get('code','')} — {r.get('title','')}", level=1)
        lines = []
        lines.append(f"Programme: {r.get('programme','-')}")
        lines.append(f"Cluster: {r.get('cluster','-')}")
        lines.append(f"Destination: {r.get('destination','-')}")
        lines.append(f"Type of Action: {r.get('type_of_action','-')}")
        trl_val = r.get("trl"); lines.append(f"TRL: {int(trl_val) if pd.notna(trl_val) else '-'}")
        ma = r.get("managing_authority"); ka = r.get("key_action")
        if pd.notna(ma) or pd.notna(ka):
            lines.append(f"Managing Authority: {ma if pd.notna(ma) else '-'}")
            lines.append(f"Key Action: {ka if pd.notna(ka) else '-'}")
        op, dl = r.get("opening_date"), r.get("deadline")
        lines.append(f"Opening: {op:%d %b %Y}" if pd.notna(op) else "Opening: -")
        lines.append(f"Deadline: {dl:%d %b %Y}" if pd.notna(dl) else "Deadline: -")
        doc.add_paragraph("\n".join(lines))
        notes = (notes_by_code or {}).get(str(r.get("code","")), "")
        doc.add_heading("Notes", level=2); doc.add_paragraph(notes if notes else "-")

    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.getvalue()

# ------------------------------------------------------------
# UI — page config + branding
page_icon = str(LOGO_SMALL) if LOGO_SMALL.exists() else "🟦"
st.set_page_config(page_title="Schuman · Funding Dashboard", page_icon=page_icon, layout="wide")
inject_brand_css()
brand_header()

st.info(
    "📂 Please upload the latest parsed Excel file.\n\n"
    "➡️ Location hint:\n\n"
    "- **3.SA Practices** → Central Systems and Bid Management → 1. Central Systems\n\n"
    "👉 Look for *Central System Funding Compass Database*.\n"
)

upl = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])
if not upl:
    st.stop()

# Detect sheets and allow override
sheets = get_sheet_names(upl.getvalue())
c1, c2 = st.columns(2)
with c1:
    hz_sheet = st.selectbox("Horizon Database", options=sheets, index=0)
with c2:
    er_sheet = st.selectbox("Erasmus Database", options=sheets, index=min(1, len(sheets)-1))

# Load each programme independently
df_h = load_programme(upl.getvalue(), hz_sheet, "Horizon Europe", _ver=1)
df_e = load_programme(upl.getvalue(), er_sheet, "Erasmus+",      _ver=1)

# ------- Build filter choices -------
all_df = pd.concat([df_h, df_e], ignore_index=True)  # programme already set
opening_years  = sorted([int(y) for y in all_df["opening_year"].dropna().unique()])
deadline_years = sorted([int(y) for y in all_df["deadline_year"].dropna().unique()])
type_opts      = sorted([t for t in all_df.get("type_of_action", pd.Series(dtype=object)).dropna().unique().tolist() if t!=""])

cluster_opts   = sorted([c for c in df_h.get("cluster", pd.Series(dtype=object)).dropna().unique().tolist() if c!=""])
dest_opts      = sorted([d for d in df_h.get("destination", pd.Series(dtype=object)).dropna().unique().tolist() if d!=""])
trl_opts       = sorted([str(int(x)) for x in df_h.get("trl", pd.Series(dtype=float)).dropna().unique() if pd.notna(x)])

ma_opts        = sorted([m for m in df_e.get("managing_authority", pd.Series(dtype=object)).dropna().unique().tolist() if m!=""])
ka_opts        = sorted([k for k in df_e.get("key_action", pd.Series(dtype=object)).dropna().unique().tolist() if k!=""])

# Budget slider (use combined range)
bud_series = pd.to_numeric(all_df.get("budget_per_project_eur"), errors="coerce").dropna()
if bud_series.empty:
    min_bud, max_bud = 0.0, 1_000_000.0
else:
    min_bud, max_bud = float(bud_series.min()), float(bud_series.max())
    if not (min_bud < max_bud):
        min_bud, max_bud = max(min_bud, 0.0), min_bud + 100000.0
rng = max_bud - min_bud
step = max(1e4, round(rng / 50, -3)) if rng else 10000.0

with st.form("filters", clear_on_submit=False):
    st.header("Filters")

    # Row 1: opening, deadline, open calls
    col_oy, col_dy, col_open = st.columns([1,1,1])
    with col_oy:
        open_year_sel = st.multiselect("Opening year(s)", opening_years, default=opening_years)
    with col_dy:
        deadline_year_sel = st.multiselect("Deadline year(s)", deadline_years, default=deadline_years)
    with col_open:
        open_calls_only = st.checkbox(
            "Open calls only",
            value=False,
            help="Keep only entries whose final Deadline is strictly after today (Europe/Brussels)."
        )

    # Row 2: budget
    col_bud, _sp = st.columns([2,1])
    with col_bud:
        budget_range = st.slider("Budget per project (EUR)", min_bud, max_bud, (min_bud, max_bud), step=step)

    # Row 3: keywords + combine + title/code toggle
    k1, k2, k3, kcomb, ktit = st.columns([2,2,2,1,1.2])
    with k1: kw1 = st.text_input("Keyword 1")
    with k2: kw2 = st.text_input("Keyword 2")
    with k3: kw3 = st.text_input("Keyword 3")
    with kcomb:
        kw_mode = st.radio("Combine", ["AND","OR"], index=1, horizontal=True)  # default OR
    with ktit:
        title_code_only = st.checkbox("Title/Code only", value=False)          # default off

    # Row 4: type of action
    types_sel = st.multiselect("Type of Action", type_opts)

    # Horizon-specific
    st.subheader("Horizon-specific")
    h1,h2,h3 = st.columns(3)
    with h1: clusters_sel = st.multiselect("Cluster", cluster_opts)
    with h2: dests_sel    = st.multiselect("Destination", dest_opts)
    with h3: trls_sel     = st.multiselect("TRL", trl_opts)

    # Erasmus-specific
    st.subheader("Erasmus+-specific")
    e1,e2 = st.columns(2)
    with e1: ma_sel = st.multiselect("Managing Authority", ma_opts)
    with e2: ka_sel = st.multiselect("Key Action", ka_opts)

    applied = st.form_submit_button("Apply filters")

# ---- track if user has applied filters at least once (welcome state)
if "has_applied" not in st.session_state:
    st.session_state.has_applied = False
if applied:
    st.session_state.has_applied = True

# Persist criteria
if "crit35" not in st.session_state:
    st.session_state.crit35 = {}
if applied or not st.session_state.crit35:
    st.session_state.crit35 = dict(
        open_years=open_year_sel, deadline_years=deadline_year_sel,
        types=types_sel, kws=[kw1,kw2,kw3], kw_mode=kw_mode, title_code_only=title_code_only,
        budget_range=budget_range,
        clusters=clusters_sel, dests=dests_sel, trls=trls_sel,
        managing_authority=ma_sel, key_action=ka_sel,
        open_calls_only=open_calls_only
    )
crit = st.session_state.crit35

# ★ shortlist state helper
def _ensure_shortlist_state():
    if "sel35" not in st.session_state: st.session_state.sel35 = set()
    if "notes35" not in st.session_state: st.session_state.notes35 = {}
    if "shortlist_chart_png" not in st.session_state: st.session_state.shortlist_chart_png = None
_ensure_shortlist_state()

# ---- performance guardrails
MAX_RENDER = 300  # tweak to taste

def guard_large_render(filt_df: pd.DataFrame, view_name: str) -> bool:
    """
    Returns True if safe to render this heavy view, else shows warning + 'Render anyway'.
    """
    n = len(filt_df)
    if n <= MAX_RENDER:
        return True
    st.warning(
        f"This {view_name} would render **{n} rows**, which may be slow. "
        f"Refine your filters or confirm to continue."
    )
    return st.button(f"Render anyway ({view_name})", key=f"render_anyway_{view_name}_{n}")

# Filtering helpers
def multi_keyword_filter(df: pd.DataFrame, terms: list[str], mode: str, title_code_only: bool) -> pd.DataFrame:
    terms = [t.strip().lower() for t in terms if t and str(t).strip()]
    if not terms:
        return df
    hay = df["_search_title"] if title_code_only else df["_search_all"]
    if mode.upper() == "AND":
        pattern = "".join([f"(?=.*{re.escape(t)})" for t in terms]) + ".*"
    else:
        pattern = "(" + "|".join(re.escape(t) for t in terms) + ")"
    return df[hay.str.contains(pattern, regex=True, na=False)]

def apply_common_filters(df0: pd.DataFrame) -> pd.DataFrame:
    df = df0.copy()
    if crit["open_years"]:
        df = df[df["opening_year"].isin(crit["open_years"])]
    if crit["deadline_years"]:
        df = df[df["deadline_year"].isin(crit["deadline_years"])]
    if crit.get("open_calls_only"):
        today = pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None)
        df = df[df["deadline"].notna() & (df["deadline"] > today)]
    if crit["types"]:
        df = df[df.get("type_of_action").isin(crit["types"])]
    lo, hi = crit["budget_range"]
    df = df[df.get("budget_per_project_eur").fillna(0).between(lo, hi)]
    df = multi_keyword_filter(df, crit["kws"], crit["kw_mode"], crit["title_code_only"])
    return df

def apply_horizon_filters(df0: pd.DataFrame) -> pd.DataFrame:
    df = apply_common_filters(df0)
    if crit["clusters"]: df = df[df.get("cluster").isin(crit["clusters"])]
    if crit["dests"]:    df = df[df.get("destination").isin(crit["dests"])]
    if crit["trls"]:
        df = df[df.get("trl").dropna().astype("Int64").astype(str).isin(crit["trls"])]
    return df

def apply_erasmus_filters(df0: pd.DataFrame) -> pd.DataFrame:
    df = apply_common_filters(df0)
    if crit["managing_authority"]: df = df[df.get("managing_authority").isin(crit["managing_authority"])]
    if crit["key_action"]:         df = df[df.get("key_action").isin(crit["key_action"])]
    return df

fh = apply_horizon_filters(df_h)
fe = apply_erasmus_filters(df_e)
st.caption(f"Rows after filters — Horizon: {len(fh)} | Erasmus: {len(fe)}")

# ---- empty state: don't render tabs until filters are applied
if not st.session_state.has_applied:
    st.markdown(
        """
        ### 👋 Welcome
        Use the **Filters** above and click **Apply filters** to load matching calls.
        - Tip: start with *Opening/Deadline year* or turn on **Open calls only**.
        """
    )
    st.stop()

# ------------------------------ Tabs ------------------------------
tab_hz, tab_er, tab_tbl, tab_full, tab_short = st.tabs([
    "📅 Gantt — Horizon", "📅 Gantt — Erasmus", "📋 Tables", "📚 Full Data", "📝 Shortlist"
])

with tab_hz:
    st.subheader("Gantt — Horizon Europe (Opening → Deadline)")
    group_by_cluster = st.checkbox(
        "Group by cluster (render one Gantt per cluster)",
        value=False,
        help="When enabled, the Horizon chart is split into one dropdown per Cluster."
    )
    if not group_by_cluster:
        if guard_large_render(fh, "Horizon Gantt"):
            g_h = build_singlebar_rows(fh)
            if g_h.empty:
                st.info("No valid Horizon rows/dates.")
            else:
                st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
                st.altair_chart(gantt_singlebar_chart(g_h, color_field="type_of_action"), use_container_width=True)
                st.markdown('</div>', unsafe_allow_html=True)
    else:
        if "cluster" not in fh.columns:
            st.warning("Column 'cluster' not available in Horizon dataset.")
        else:
            tmp = fh.copy()
            tmp["cluster"] = tmp["cluster"].fillna("— Unspecified —").replace({"": "— Unspecified —"})
            groups = list(tmp.groupby("cluster", dropna=False))
            groups.sort(key=lambda kv: len(kv[1]), reverse=True)
            st.caption(f"Clusters found: {len(groups)}")
            for clu_name, gdf in groups:
                if not guard_large_render(gdf, f"Horizon Gantt · {clu_name}"):
                    continue
                g_clu = build_singlebar_rows(gdf)
                with st.expander(f"Cluster: {clu_name}  ({len(g_clu)} calls)", expanded=False):
                    if g_clu.empty:
                        st.info("No valid rows/dates for this cluster after filters.")
                    else:
                        st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
                        st.altair_chart(gantt_singlebar_chart(g_clu, color_field="type_of_action"), use_container_width=True)
                        st.markdown('</div>', unsafe_allow_html=True)

with tab_er:
    st.subheader("Gantt — Erasmus+ (Opening → Deadline)")
    if guard_large_render(fe, "Erasmus Gantt"):
        g_e = build_singlebar_rows(fe)
        if g_e.empty: st.info("No valid Erasmus rows/dates.")
        else:
            st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
            st.altair_chart(gantt_singlebar_chart(g_e, color_field="type_of_action"), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

with tab_tbl:
    st.subheader("Tables")

    def show_paged_df(df: pd.DataFrame, label: str, cols: list[str], page_size_default: int = 50):
        if df.empty:
            st.caption("— no rows —")
            return
        with st.expander(f"{label} — options", expanded=False):
            page_size = st.number_input("Rows per page", 10, 500, page_size_default, 10, key=f"ps_{label}")
            max_page = max(1, (len(df)-1)//page_size + 1)
            page = st.number_input("Page", 1, max_page, 1, 1, key=f"pg_{label}")
        start = (page-1)*page_size
        end = start + page_size
        st.dataframe(df[cols].iloc[start:end], use_container_width=True, hide_index=True)
        st.caption(f"Showing {min(end, len(df))}/{len(df)}")

    show_cols_h = [c for c in DISPLAY_COLS if c in fh.columns]
    show_cols_e = [c for c in DISPLAY_COLS if c in fe.columns]

    st.markdown(f"### Horizon Europe ({len(fh)})")
    show_paged_df(fh, "Horizon", show_cols_h)

    st.markdown(f"### Erasmus+ ({len(fe)})")
    show_paged_df(fe, "Erasmus", show_cols_e)

# helper to render shortlist checkbox + row expander as a single “row”
def render_shortlist_row(exp_label: str, code: str, render_body_fn):
    row_cols = st.columns([0.16, 0.84])
    with row_cols[0]:
        checked = code in st.session_state.sel35
        new_val = st.checkbox("Shortlist", value=checked, key=f"short_{code}_{hash(exp_label) % 10_000_000}")
        if new_val and not checked:
            st.session_state.sel35.add(code)
        elif not new_val and checked:
            st.session_state.sel35.discard(code)
    with row_cols[1]:
        with st.expander(exp_label or "(untitled)"):
            render_body_fn()

with tab_full:
    st.subheader("Full Data — Expand rows for details")

    kw_list = crit.get("kws", [])

    def render_row(row, programme: str):
        c1, c2 = st.columns(2)
        with c1:
            if pd.notna(row.get("opening_date")):
                st.markdown(f"📅 **Opening:** {row['opening_date']:%d %b %Y}")
            if pd.notna(row.get("deadline")):
                st.markdown(f"⏳ **Deadline:** {row['deadline']:%d %b %Y}")
        with c2:
            if pd.notna(row.get("budget_per_project_eur")):
                st.markdown(f"💶 **Budget/Project:** {row['budget_per_project_eur']:,.0f} EUR")
            if pd.notna(row.get("total_budget_eur")):
                st.markdown(f"📦 **Total:** {row['total_budget_eur']:,.0f} EUR")
            if pd.notna(row.get("num_projects")):
                st.markdown(f"📊 **# Projects:** {int(row['num_projects'])}")

        meta_bits = [
            f"🏷️ **Programme:** {programme}",
            f"**Type of Action:** {row.get('type_of_action','-')}",
        ]
        if programme == "Horizon Europe":
            meta_bits += [
                f"**Cluster:** {row.get('cluster','-')}",
                f"**Destination:** {row.get('destination','-')}",
                f"**TRL:** {row.get('trl','-')}",
            ]
        else:
            meta_bits += [
                f"**Managing Authority:** {row.get('managing_authority','-')}",
                f"**Key Action:** {row.get('key_action','-')}",
            ]
        st.markdown(" | ".join(meta_bits))

        if row.get("expected_outcome"):
            with st.expander("🎯 Expected Outcome"):
                clean_text = nl_to_br(normalize_bullets(clean_footer(str(row.get("expected_outcome")))))
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)
        if row.get("scope"):
            with st.expander("🧭 Scope"):
                clean_text = nl_to_br(normalize_bullets(clean_footer(str(row.get("scope")))))
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)
        if row.get("full_text"):
            with st.expander("📖 Full Description"):
                clean_text = nl_to_br(normalize_bullets(clean_footer(str(row.get("full_text")))))
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)

        st.caption(
            f"📂 Source: {row.get('source_filename','-')} "
            f"| Version: {row.get('version_label','-')} "
            f"| Parsed on: {row.get('parsed_on_utc','-')}"
        )

    # ----- HORIZON -----
    st.markdown(f"### Horizon Europe ({len(fh)})")
    if not guard_large_render(fh, "Full Data — Horizon"):
        st.stop()
    group_hz = st.checkbox(
        "Group Horizon by Cluster (dropdowns)",
        value=False,
        help="Show one expander per Cluster; inside each, expand rows for details."
    )
    if len(fh) == 0:
        st.caption("— no Horizon rows after filters —")
    else:
        if not group_hz:
            for i, r in fh.iterrows():
                label = f"{str(r.get('code') or '')} — {str(r.get('title') or '')}".strip(" —")
                code = str(r.get("code") or f"id-{i}")
                render_shortlist_row(
                    label, code,
                    lambda rr=r: render_row(rr, "Horizon Europe")
                )
        else:
            tmp = fh.copy()
            tmp["cluster"] = tmp.get("cluster").fillna("— Unspecified —").replace({"": "— Unspecified —"})
            groups = list(tmp.groupby("cluster", dropna=False))
            groups.sort(key=lambda kv: len(kv[1]), reverse=True)
            st.caption(f"Clusters found: {len(groups)}")
            for clu_name, gdf in groups:
                disp = str(clu_name)
                with st.expander(f"Cluster: {disp}  ({len(gdf)} calls)", expanded=False):
                    for i, r in gdf.iterrows():
                        label = f"{str(r.get('code') or '')} — {str(r.get('title') or '')}".strip(" —")
                        code = str(r.get("code") or f"id-{i}")
                        render_shortlist_row(
                            label, code,
                            lambda rr=r: render_row(rr, "Horizon Europe")
                        )

    # ----- ERASMUS -----
    st.markdown(f"### Erasmus+ ({len(fe)})")
    if not guard_large_render(fe, "Full Data — Erasmus"):
        st.stop()
    if len(fe) == 0:
        st.caption("— no Erasmus rows after filters —")
    else:
        for i, r in fe.iterrows():
            label = f"{str(r.get('code') or '')} — {str(r.get('title') or '')}".strip(" —")
            code = str(r.get("code") or f"id-{i}")
            render_shortlist_row(
                label, code,
                lambda rr=r: render_row(rr, "Erasmus+")
            )

with tab_short:
    st.subheader("Shortlist & Notes (DOCX)")
    _ensure_shortlist_state()

    # Build combined filtered set
    combined = []
    if len(fh): combined.append(fh)
    if len(fe): combined.append(fe)
    ff = pd.concat(combined, ignore_index=True) if combined else pd.DataFrame()

    # Derive shortlist view from session state
    shortlisted_codes = set(st.session_state.sel35)
    if ff.empty or not shortlisted_codes:
        st.info("No shortlisted calls yet. Use the **Shortlist** checkbox in the Full Data tab to add items.")
        st.stop()

    # Ensure closing_date_any exists (should already be present)
    if "closing_date_any" not in ff.columns:
        close_cols = [c for c in ["deadline","first_deadline","second_deadline"] if c in ff.columns]
        if close_cols:
            ff["closing_date_any"] = pd.to_datetime(ff[close_cols].stack(), errors="coerce").groupby(level=0).min()
        else:
            ff["closing_date_any"] = pd.NaT

    ff = ff.sort_values([c for c in ["closing_date_any","opening_date"] if c in ff.columns])
    selected_df = ff[ff["code"].astype(str).isin(shortlisted_codes)]

    st.markdown("**Your shortlisted calls** (you can uncheck to remove):")
    for idx, r in selected_df.iterrows():
        code = str(r.get("code") or "")
        title = str(r.get("title") or "")
        label = f"{code} — {title}".strip(" —")
        checked = code in st.session_state.sel35
        new_val = st.checkbox(label or "(untitled)", value=checked, key=f"sel35_mirror_{code}_{idx}")
        if new_val and not checked: st.session_state.sel35.add(code)
        elif (not new_val) and checked: st.session_state.sel35.discard(code)

    # Recompute selected set after potential toggles
    shortlisted_codes = set(st.session_state.sel35)
    selected_df = ff[ff["code"].astype(str).isin(shortlisted_codes)]

    # --- Gantt for shortlisted items (Altair in-app, Matplotlib for DOCX) ---
    if guard_large_render(selected_df, "Shortlist Gantt"):
        st.markdown("### 📅 Gantt — Shortlisted Calls")

        gsrc = _prepare_dates_for_chart(selected_df)

        col_g1, col_g2 = st.columns([1,1])
        with col_g1:
            color_by = st.selectbox(
                "Colour bars by",
                options=[opt for opt in ["type_of_action", "programme", "cluster"] if opt in gsrc.columns],
                index=0
            )
        with col_g2:
            group_hz_clusters = st.checkbox(
                "For Horizon: split into one chart per Cluster",
                value=False
            )

        g_all = build_singlebar_rows(gsrc)
        if g_all.empty:
            total = len(selected_df)
            with_dates = selected_df["deadline"].notna().sum()
            st.info(f"No valid date ranges for the current shortlist. (Selected: {total}; with any deadline: {with_dates})")
            st.session_state.shortlist_chart_png = None
        else:
            st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
            st.altair_chart(gantt_singlebar_chart(g_all, color_field=color_by), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
            # Build the PNG once for DOCX
            st.session_state.shortlist_chart_png = _shortlist_gantt_png(gsrc, color_by=color_by)

        # Optional: one Gantt per Horizon cluster (display-only)
        if group_hz_clusters and not g_all.empty:
            hz_only = gsrc[gsrc.get("programme").eq("Horizon Europe")]
            if not hz_only.empty and "cluster" in hz_only.columns:
                tmp = hz_only.copy()
                tmp["cluster"] = tmp["cluster"].fillna("— Unspecified —").replace({"": "— Unspecified —"})
                groups = list(tmp.groupby("cluster", dropna=False))
                groups.sort(key=lambda kv: len(kv[1]), reverse=True)
                st.caption(f"Horizon clusters in shortlist: {len(groups)}")
                for clu_name, gdf in groups:
                    g_clu = build_singlebar_rows(gdf)
                    with st.expander(f"Cluster: {clu_name}  ({len(g_clu)} calls)", expanded=False):
                        if g_clu.empty:
                            st.info("No valid rows/dates for this cluster.")
                        else:
                            st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
                            st.altair_chart(gantt_singlebar_chart(g_clu, color_field=color_by), use_container_width=True)
                            st.markdown('</div>', unsafe_allow_html=True)

    # Notes + DOCX
    if not selected_df.empty:
        st.markdown("---")
        for _, r in selected_df.iterrows():
            code = str(r.get("code") or "")
            default = st.session_state.notes35.get(code, "")
            st.session_state.notes35[code] = st.text_area(f"Notes — {code}", value=default, height=110, key=f"note35_{code}")

        colA, _colB = st.columns(2)
        with colA: title = st.text_input("Report title", value="Funding Report – Shortlist (app_b4.1)")

        if st.button("📄 Generate DOCX"):
            try:
                if DOCX_AVAILABLE:
                    data = generate_docx_report(
                        selected_df,
                        st.session_state.notes35,
                        title=title,
                        shortlist_gantt_png=st.session_state.shortlist_chart_png
                    )
                    st.download_button("⬇️ Download .docx", data=data,
                                       file_name="funding_report.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error("python-docx not installed in this environment.")
            except Exception as e:
                st.error(f"Failed to generate report: {e}")
    else:
        st.info("Your shortlist is now empty.")
