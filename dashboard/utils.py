# utils.py — helpers (branding, loading, cleaning, charts, shortlist, report)
from __future__ import annotations
import io, re, base64
from datetime import datetime
from typing import List, Dict, Optional, Iterable
from pathlib import Path

import pandas as pd
import streamlit as st
import altair as alt

# DOCX
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Matplotlib for PNG chart embedding
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates

# ---------- Brand assets ----------
ASSETS_DIR = Path("assets")
FONTS_DIR  = ASSETS_DIR / "fonts"
LOGO_BLUE  = ASSETS_DIR / "logo-schuman_blue.png"
LOGO_GREY  = ASSETS_DIR / "logo-schuman_grey.png"
LOGO_WHITE = ASSETS_DIR / "logo-schuman_white.png"

FONT_FILES = [
    ("SA Brand", "normal", 300, ASSETS_DIR / "Frm-Light.otf"),
    ("SA Brand", "normal", 400, ASSETS_DIR / "Frm-Regular.otf"),
    ("SA Brand", "normal", 500, ASSETS_DIR / "Frm-Medium.otf"),
    ("SA Brand", "normal", 700, ASSETS_DIR / "Frm-Bold.otf"),
    ("SA Brand", "italic", 700, ASSETS_DIR / "Frm-Bold-Italic.otf"),
    ("SA Brand", "normal", 900, ASSETS_DIR / "Frm-Black.otf"),
]

# ---------- Columns / display ----------
COLUMN_MAP = {
    "Programme": "programme", "Code": "code", "Title": "title",
    "Opening Date": "opening_date", "Opening date": "opening_date",
    "Deadline": "deadline", "First Stage Deadline": "first_deadline",
    "Second Stage Deadline": "second_deadline", "Second Stage deadline": "second_deadline",
    "Two-Stage": "two_stage", "Cluster": "cluster", "Destination": "destination",
    "Destination / Strand": "destination", "Destination/Strand": "destination", "Strand": "destination",
    "Budget Per Project": "budget_per_project_eur", "Budget per project": "budget_per_project_eur",
    "Total Budget": "total_budget_eur", "Number of Projects": "num_projects",
    "Type of Action": "type_of_action", "TRL": "trl", "Call Name": "call_name",
    "Expected Outcome": "expected_outcome", "Scope": "scope", "Description": "full_text",
    "Source Filename": "source_filename", "Version Label": "version_label", "Parsed On (UTC)": "parsed_on_utc",
    "Managing Authority": "managing_authority", "Key Action": "key_action",
}
DISPLAY_COLS = [
    "programme","code","title","opening_date","deadline",
    "type_of_action","budget_per_project_eur",
    "cluster","destination","trl",
    "managing_authority","key_action",
    "first_deadline","second_deadline","two_stage",
    "call_name","version_label","source_filename","parsed_on_utc",
]

# ---------- General small utils ----------
@st.cache_data(show_spinner=False)
def _file_to_base64(p: Path) -> str | None:
    try:
        return base64.b64encode(p.read_bytes()).decode("utf-8")
    except Exception:
        return None

@st.cache_data(show_spinner=False)
def _read_theme_css(path: Path) -> str:
    _ = path.stat().st_mtime if path.exists() else 0
    return path.read_text(encoding="utf-8") if path.exists() else ""

def inject_brand_css():
    """Injects @font-face (from OTFs) + static theme from assets/theme.css."""
    font_faces = []
    for fam, style, weight, path in FONT_FILES:
        if path.exists():
            b64 = _file_to_base64(path)
            if b64:
                font_faces.append(f"""
                @font-face {{
                  font-family: '{fam}';
                  src: url(data:font/otf;base64,{b64}) format('opentype');
                  font-style: {style};
                  font-weight: {weight};
                  font-display: swap;
                }}
                """)
    font_css = "\n".join(font_faces)
    theme_css = _read_theme_css(ASSETS_DIR / "theme.css")
    st.markdown(f"<style>\n{font_css}\n{theme_css}\n</style>", unsafe_allow_html=True)

def brand_header():
    """Top hero header with brand blue and logo."""
    logo_src = None
    for p in (LOGO_WHITE, LOGO_GREY, LOGO_BLUE):
        if p.exists():
            logo_src = f"data:image/png;base64,{_file_to_base64(p)}"
            break
    st.markdown(f"""
    <div style="
      border-radius: 16px;
      background: var(--sa-primary, #1E4F86);
      padding: 24px 20px;
      color: white;
      text-align: center;
      box-shadow: var(--sa-shadow, 0 6px 24px rgba(23,32,84,0.08));">
      {'<img src="'+logo_src+'" alt="Schuman Associates" style="height:60px; margin-bottom:12px;" />' if logo_src else ''}
      <div style="font-size:20px; font-weight:700; margin-bottom:4px;">Schuman Associates · Funding Dashboard</div>
      <div style="font-size:14px; opacity:.9;">Your European partners in a global market since 1989</div>
    </div>
    """, unsafe_allow_html=True)

# ---------- Text utils ----------

def nl_to_br(s: str) -> str:
    """
    Replace newline characters with <br> for safe HTML rendering.
    """
    return "" if not s else s.replace("\n", "<br>")


def highlight_text(text: str, keywords: list[str], colours=None, match_case: bool = False) -> str:
    """
    Wrap keywords in <span> with background colour + bold.
    - text: input string
    - keywords: list of terms to highlight
    - colours: list of colours to cycle through
    - match_case: if True, highlight case-sensitively; otherwise case-insensitive
    """
    import re

    if not text:
        return ""

    kws = [str(k).strip() for k in keywords if k and str(k).strip()]
    if not kws:
        return text

    if colours is None:
        colours = ["#ffff00", "#a0e7e5", "#ffb3b3"]

    out = str(text)
    for i, kw in enumerate(kws):
        colour = colours[i % len(colours)]
        flags = 0 if match_case else re.IGNORECASE
        out = re.sub(
            re.escape(kw),
            lambda m: f"<span style='background-color:{colour}; font-weight:bold;'>{m.group(0)}</span>",
            out,
            flags=flags,
        )
    return out


def merge_edits_into_df(df: pd.DataFrame, sstate) -> None:
    """In-place: apply text edits from session_state to df, if present."""
    for i, row in df.iterrows():
        code = str(row.get("code") or f"id-{i}")
        for field in ("expected_outcome", "scope", "full_text"):
            k = f"edit_{field}_{code}"
            if k in sstate and isinstance(sstate[k], str) and sstate[k].strip():
                df.at[i, field] = sstate[k]


def clean_footer(text: str) -> str:
    if not text:
        return ""
    pat = re.compile(r"Horizon\s*Europe\s*[-–]?\s*Work Programme.*?Page\s+\d+\s+of\s+\d+", re.IGNORECASE | re.DOTALL)
    cleaned = pat.sub("", text)
    return re.sub(r"\s+", " ", cleaned).strip()

def normalize_bullets(text: str) -> str:
    """
    Stub: do nothing, return text unchanged.
    """
    if not isinstance(text, str):
        return ""
    return text


def strip_and_collect_footnotes(text: str) -> tuple[str, dict[int, str]]:
    """Neutered placeholder — returns text unchanged and empty dict."""
    if not isinstance(text, str) or not text.strip():
        return "", {}
    return text, {}

# ---------- Dates & canonicalisation ----------
def safe_date_series(s: pd.Series) -> pd.Series:
    out = pd.to_datetime(s, errors="coerce", dayfirst=True)
    if out.notna().sum() == 0:
        out = pd.to_datetime(s, errors="coerce", dayfirst=False)
    return out

def canonicalise(df: pd.DataFrame, programme_name: str) -> pd.DataFrame:
    df.columns = [c.strip() for c in df.columns]
    for src, dst in COLUMN_MAP.items():
        if src in df.columns and dst not in df.columns:
            df = df.rename(columns={src: dst})
    df = df.rename(columns={c: c.strip().lower() for c in df.columns})
    df["programme"] = programme_name

    for c in ("budget_per_project_eur","total_budget_eur","trl","num_projects"):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    for c in ("opening_date","deadline","first_deadline","second_deadline"):
        if c in df.columns:
            df[c] = safe_date_series(df[c])

    if "two_stage" in df.columns:
        df["two_stage"] = (
            df["two_stage"].astype(str).str.strip().str.lower()
            .map({"true": True, "false": False, "yes": True, "no": False, "1": True, "0": False})
            .fillna(False)
        )
    else:
        df["two_stage"] = False

    # --- searchable fields ---
    present = [c for c in (
        "code","title","call_name","expected_outcome","scope","full_text",
        "cluster","destination","type_of_action","trl","managing_authority","key_action"
    ) if c in df.columns]

    if present:
        joined = df[present].astype(str).agg(" ".join, axis=1)
        df["_search_all"] = joined.str.lower()
        df["_search_all_raw"] = joined
    else:
        df["_search_all"] = ""
        df["_search_all_raw"] = ""   # ensure always present

    title_cols = [c for c in ["code","title"] if c in df.columns]
    if title_cols:
        joined_titles = df[title_cols].astype(str).agg(" ".join, axis=1)
        df["_search_title"] = joined_titles.str.lower()
        df["_search_title_raw"] = joined_titles
    else:
        df["_search_title"] = ""
        df["_search_title_raw"] = ""   # ensure always present

    # --- dates ---
    close_cols = [c for c in ["deadline","first_deadline","second_deadline"] if c in df.columns]
    if close_cols:
        df["closing_date_any"] = pd.to_datetime(df[close_cols].stack(), errors="coerce").groupby(level=0).min()
    else:
        df["closing_date_any"] = pd.NaT
    df["opening_year"]  = df["opening_date"].dt.year
    df["deadline_year"] = df["deadline"].dt.year

    return df

# ---------- Caching: load sheets ----------
@st.cache_data(show_spinner=False)
def get_sheet_names(file_bytes: bytes) -> List[str]:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    return xls.sheet_names

@st.cache_data(show_spinner=False)
def load_programme(file_bytes: bytes, sheet_name: str, programme_name: str, _ver:int=1) -> pd.DataFrame:
    _ = hash(file_bytes)
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    raw = pd.read_excel(xls, sheet_name=sheet_name)
    df = canonicalise(raw, programme_name)
    return df.copy(deep=True)

# ---------- Charts (Altair) ----------
def wrap_label(text: str, width=60, max_lines=3) -> str:
    s = str(text or "")
    chunks = [s[i:i+width] for i in range(0, len(s), width)]
    return "\n".join(chunks[:max_lines]) if chunks else "—"

def build_month_bands(min_x: pd.Timestamp, max_x: pd.Timestamp) -> pd.DataFrame:
    start = pd.Timestamp(min_x).to_period("M").start_time
    end   = (pd.Timestamp(max_x).to_period("M") + 1).start_time
    months = pd.date_range(start, end, freq="MS")
    return pd.DataFrame({"start": months[:-1], "end": months[1:], "band": [i % 2 for i in range(len(months)-1)]})

def build_singlebar_rows(df: pd.DataFrame) -> pd.DataFrame:
    g = df.copy()
    if "code" in g.columns and g["code"].notna().any():
        base = g["code"].fillna("").astype(str)
    elif "title" in g.columns and g["title"].notna().any():
        base = g["title"].fillna("").astype(str)
    else:
        base = pd.Series([f"row-{i}" for i in range(len(g))], index=g.index)
    if base.duplicated(keep=False).any():
        base = base + g.groupby(base).cumcount().astype(str).radd("#")

    g["y_label"] = base.map(lambda s: wrap_label(s, width=100, max_lines=5))
    g["title_inbar"] = g.get("title","").astype(str).map(lambda s: wrap_label(s, width=100, max_lines=3))
    g = g[pd.notna(g["opening_date"]) & pd.notna(g["deadline"]) & (g["opening_date"] <= g["deadline"])].copy()
    if g.empty:
        return g
    g["bar_days"] = (g["deadline"] - g["opening_date"]).dt.days
    g["mid"] = g["opening_date"] + (g["deadline"] - g["opening_date"])/2
    return g.sort_values(["deadline","opening_date"])

def sa_altair_theme():
    return {
        "config": {
            "font": "SA Brand",
            "axis":   {"labelFont": "SA Brand", "labelFontWeight": 500, "titleFont": "SA Brand", "labelColor":"#0F172A"},
            "legend": {"labelFont": "SA Brand", "titleFont": "SA Brand"},
            "header": {"labelFont": "SA Brand", "titleFont": "SA Brand"},
            "title":  {"font": "SA Brand", "fontSize": 16, "fontWeight": 700, "color":"#0F172A"},
            "range": {
                "category": ["#1E4F86","#66C2A5","#FC8D62","#8DA0CB","#E78AC3","#A6D854","#FFD92F","#E5C494","#B3B3B3"]
            },
            "view": {"stroke": "transparent"}
        }
    }

alt.themes.register("sa_theme", sa_altair_theme)
alt.themes.enable("sa_theme")

def gantt_singlebar_chart(g: pd.DataFrame, color_field: str = "type_of_action", title: str = ""):
    if g is None or g.empty:
        return None

    min_x = min(g["opening_date"].min(), g["deadline"].min())
    max_x = max(g["opening_date"].max(), g["deadline"].max())
    bands_df = build_month_bands(min_x, max_x)

    month_shade = alt.Chart(bands_df).mark_rect(tooltip=False).encode(
        x="start:T", x2="end:T",
        opacity=alt.Opacity("band:Q", scale=alt.Scale(domain=[0,1], range=[0.0,0.05]), legend=None),
        color=alt.value("#1E4F86")
    )
    months = pd.date_range(pd.Timestamp(min_x).to_period("M").start_time,
                           pd.Timestamp(max_x).to_period("M").end_time, freq="MS")
    month_grid = alt.Chart(pd.DataFrame({"t": months})).mark_rule(stroke="#1E4F86", strokeWidth=0.3).encode(x="t:T")
    month_labels_df = pd.DataFrame({
        "month": months[:-1], "next_month": months[1:],
        "label": [m.strftime("%b %Y") for m in months[:-1]]
    })
    month_labels_df["mid"] = month_labels_df["month"] + ((month_labels_df["next_month"] - month_labels_df["month"]) / 2)
    month_labels = alt.Chart(month_labels_df).mark_text(
        align="center", baseline="top", dy=0, fontSize=11, fontWeight="bold"
    ).encode(x="mid:T", text="label:N", y=alt.value(0))

    today_ts = pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None)
    today_df = pd.DataFrame({"t":[today_ts]})
    today_rule = alt.Chart(today_df).mark_rule(color="#1E4F86", strokeDash=[2,1], strokeWidth=2).encode(
        x="t:T", tooltip=[alt.Tooltip("t:T", title="Today", format="%d %b %Y")]
    )
    today_label = alt.Chart(today_df).mark_text(
        align="left", baseline="top", dx=4, dy=18, fontSize=11, fontWeight="bold", color="#1E4F86"
    ).encode(x="t:T", y=alt.value(0), text=alt.Text("t:T", format='Today: "%d %b %Y"'))

    y_order = g["y_label"].drop_duplicates().tolist()
    row_h = 46
    bar_size = int(row_h * 0.38)
    domain_min, domain_max = g["opening_date"].min(), g["deadline"].max()

    base = alt.Chart(g).encode(
        y=alt.Y("y_label:N", sort=y_order,
                axis=alt.Axis(title=None, labelLimit=200, labelFontSize=11, labelAlign="right", labelPadding=50, domain=True),
                scale=alt.Scale(domain=y_order, paddingInner=0.3, paddingOuter=0.8))
    )

    bars = base.mark_bar(cornerRadius=10, size=bar_size).encode(
        x=alt.X("opening_date:T",
                axis=alt.Axis(title=None, format="%b %Y", tickCount="month", orient="top",
                              labelFontSize=11, labelPadding=50, labelOverlap="greedy", tickSize=6),
                scale=alt.Scale(domain=[domain_min, domain_max])),
        x2="deadline:T",
        color=alt.Color(color_field + ":N",
                        legend=alt.Legend(title=color_field.replace("_"," ").title(),
                                          orient="top", direction="horizontal", offset=100),
                        scale=alt.Scale(scheme="set2")),
        tooltip=[
            alt.Tooltip("title:N", title="Title"),
            alt.Tooltip("opening_date:T", title="Opening", format="%d %b %Y"),
            alt.Tooltip("deadline:T", title="Deadline", format="%d %b %Y"),
            alt.Tooltip(color_field + ":N", title=color_field.replace("_"," ").title()),
        ]
    )

    start_labels = base.mark_text(align="right", dx=-4, dy=5, fontSize=10, color="#111")\
        .encode(x="opening_date:T", text=alt.Text("opening_date:T", format="%d %b %Y"))
    end_labels   = base.mark_text(align="left",  dx=4, dy=5, fontSize=10, color="#111")\
        .encode(x="deadline:T",      text=alt.Text("deadline:T",      format="%d %b %Y"))
    
    inbar = base.mark_text(align="left",
                           baseline="bottom",
                           dx=2, 
                           dy=-(int(bar_size/2)+4),
                           color="black",
                           font="SA Brand",
                           fontSize=12,
                           fontWeight=700)\
        .encode(x=alt.X("opening_date:T",
                        scale=alt.Scale(domain=[domain_min, domain_max]),
                        axis=None),
                text="title_inbar:N",
                opacity=alt.condition(alt.datum.bar_days >= 10, alt.value(1), alt.value(0)))

    chart = (month_shade + month_grid + bars + start_labels + end_labels + inbar + month_labels + today_rule + today_label)\
        .properties(height=max(800, len(y_order)*row_h), width='container',
                    padding={"top":50,"bottom":30,"left":10,"right":10})\
        .configure_axis(grid=False, domain=True, domainWidth=1)\
        .configure_view(continuousHeight=500, continuousWidth=500, strokeWidth=0, clip=False)\
        .interactive(bind_x=True)

    return chart if not title else chart.properties(title=title)

# ---------- Charts (matplotlib PNG for DOCX) ----------
def prepare_dates_for_chart(df: pd.DataFrame, default_window_days: int = 60) -> pd.DataFrame:
    g = df.copy()
    end = g.get("deadline")
    if "closing_date_any" in g.columns:
        end = end.fillna(g["closing_date_any"]) if end is not None else g["closing_date_any"]
    g["_chart_end"] = end
    start = g.get("opening_date")
    g["_chart_start"] = start
    have_end_only = g["_chart_start"].isna() & g["_chart_end"].notna()
    g.loc[have_end_only, "_chart_start"] = g.loc[have_end_only, "_chart_end"] - pd.Timedelta(days=default_window_days)
    g = g[g["_chart_start"].notna() & g["_chart_end"].notna()].copy()
    inverted = g["_chart_start"] > g["_chart_end"]
    g.loc[inverted, "_chart_start"] = g.loc[inverted, "_chart_end"] - pd.Timedelta(days=7)
    g["opening_date"] = g["_chart_start"]
    g["deadline"]     = g["_chart_end"]
    return g

def shortlist_gantt_png(df: pd.DataFrame, color_by: str = "type_of_action") -> Optional[bytes]:
    if df is None or df.empty:
        return None
    g = prepare_dates_for_chart(df)
    g = g[pd.notna(g["opening_date"]) & pd.notna(g["deadline"]) & (g["opening_date"] <= g["deadline"])].copy()
    if g.empty:
        return None

    base = g["code"].fillna("").astype(str)
    fallback = g["title"].fillna("").astype(str)
    labels = base.where(base.ne(""), fallback)
    labels = labels + g.groupby(labels).cumcount().replace(0, "").astype(str).radd("#").replace("#0", "")
    g["_y"] = labels
    g = g.sort_values(["deadline","opening_date"]).reset_index(drop=True)

    categories = g.get(color_by) if color_by in g.columns else None
    if categories is not None:
        cats = categories.fillna("—").astype(str)
        uniq = pd.unique(cats)
        cmap = plt.get_cmap("tab20")
        color_map = {c: cmap(i % 20) for i, c in enumerate(uniq)}
        colors = cats.map(color_map)
    else:
        colors = None

    n = len(g)
    h = max(3.0, min(0.4 * n + 1.0, 10.0))
    fig, ax = plt.subplots(figsize=(10, h), dpi=160)

    start_nums = mdates.date2num(pd.to_datetime(g["opening_date"]).dt.to_pydatetime())
    dur = (pd.to_datetime(g["deadline"]) - pd.to_datetime(g["opening_date"])).dt.days.clip(lower=1)

    y_pos = range(n)
    ax.barh(list(y_pos), dur, left=start_nums, height=0.35,
            align="center", color=(colors if colors is not None else None), edgecolor="none")

    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(g["_y"])
    ax.invert_yaxis()

    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))

    today = mdates.date2num(pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None).to_pydatetime())
    ax.axvline(today, linestyle="--", linewidth=1.5, color="#1E4F86")

    ax.grid(axis="x", linestyle=":", linewidth=0.5, alpha=0.6)
    ax.set_xlabel("")
    ax.set_ylabel("")
    plt.tight_layout()

    bio = io.BytesIO()
    plt.savefig(bio, format="png", bbox_inches="tight")
    plt.close(fig)
    bio.seek(0)
    return bio.getvalue()

# ---------- Shortlist & state ----------
MAX_RENDER = 300

def ensure_shortlist_state():
    if "sel35" not in st.session_state: st.session_state.sel35 = set()
    if "notes35" not in st.session_state: st.session_state.notes35 = {}

def guard_large_render(filt_df: pd.DataFrame, view_name: str) -> bool:
    n = len(filt_df)
    if n <= MAX_RENDER:
        return True
    st.warning(
        f"This {view_name} would render **{n} rows**, which may be slow. "
        f"Refine your filters or confirm to continue."
    )
    return st.button(f"Render anyway ({view_name})", key=f"render_anyway_{view_name}_{n}")

def render_shortlist_row(exp_label: str, code: str, render_body_fn):
    row_cols = st.columns([0.16, 0.84])
    with row_cols[0]:
        checked = code in st.session_state.sel35
        new_val = st.checkbox("Shortlist", value=checked, key=f"short_{code}_{hash(exp_label) % 10_000_000}")
        if new_val and not checked:
            st.session_state.sel35.add(code)
        elif not new_val and checked:
            st.session_state.sel35.discard(code)
    with row_cols[1]:
        with st.expander(exp_label or "(untitled)"):
            render_body_fn()

def merge_edits_into_df(df: pd.DataFrame, sstate) -> None:
    """In-place: apply text edits from session_state to df."""
    if df.empty:
        return
    for idx, rr in df.iterrows():
        code = str(rr.get("code") or "")
        for field in ["expected_outcome","scope","full_text"]:
            key = f"edit_{field}_{code}"
            val = sstate.get(key, "")
            if isinstance(val, str) and val.strip():
                df.at[idx, field] = val

# ---------- Report ----------
def generate_docx_report(calls_df: pd.DataFrame, notes_by_code: Dict[str,str], title="Funding Report",
                         shortlist_gantt_png: Optional[bytes] = None) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    h = doc.add_heading(title, level=0); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"Generated on {datetime.utcnow():%d %b %Y, %H:%M UTC}"); p.runs[0].font.size = Pt(9)

    # Include shortlist Gantt image if provided
    if shortlist_gantt_png:
        doc.add_heading("Shortlist Gantt", level=1)
        doc.add_picture(io.BytesIO(shortlist_gantt_png), width=Inches(6.5))

    table = doc.add_table(rows=1, cols=5); hdr = table.rows[0].cells
    for i, t in enumerate(["Programme","Code","Title","Opening","Deadline"]): hdr[i].text = t
    for _, r in calls_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r.get("programme","")); row[1].text = str(r.get("code","")); row[2].text = str(r.get("title",""))
        op, dl = r.get("opening_date"), r.get("deadline")
        row[3].text = op.strftime("%d %b %Y") if pd.notna(op) else "-"
        row[4].text = dl.strftime("%d %b %Y") if pd.notna(dl) else "-"

    for _, r in calls_df.iterrows():
        doc.add_page_break()
        doc.add_heading(f"{r.get('code','')} — {r.get('title','')}", level=1)
        lines = []
        lines.append(f"Programme: {r.get('programme','-')}")
        lines.append(f"Cluster: {r.get('cluster','-')}")
        lines.append(f"Destination: {r.get('destination','-')}")
        lines.append(f"Type of Action: {r.get('type_of_action','-')}")
        trl_val = r.get("trl"); lines.append(f"TRL: {int(trl_val) if pd.notna(trl_val) else '-'}")
        ma = r.get("managing_authority"); ka = r.get("key_action")
        if pd.notna(ma) or pd.notna(ka):
            lines.append(f"Managing Authority: {ma if pd.notna(ma) else '-'}")
            lines.append(f"Key Action: {ka if pd.notna(ka) else '-'}")
        op, dl = r.get("opening_date"), r.get("deadline")
        lines.append(f"Opening: {op:%d %b %Y}" if pd.notna(op) else "Opening: -")
        lines.append(f"Deadline: {dl:%d %b %Y}" if pd.notna(dl) else "Deadline: -")
        doc.add_paragraph("\n".join(lines))

        # Include edited texts if any
        for field, label in [
            ("expected_outcome","Expected Outcome"),
            ("scope","Scope"),
            ("full_text","Full Description"),
        ]:
            val = str(r.get(field) or "").strip()
            if val:
                doc.add_heading(label, level=2)
                doc.add_paragraph(val)

        notes = (notes_by_code or {}).get(str(r.get("code","")), "")
        doc.add_heading("Notes", level=2); doc.add_paragraph(notes if notes else "-")

    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.getvalue()
