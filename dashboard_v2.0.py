# app_b.py — Streamlit Funding Dashboard (optimised) + Shortlist Notes → DOCX/HTML
from __future__ import annotations

import io
import base64
import re
from datetime import datetime
from typing import List, Dict

import pandas as pd
import streamlit as st
import altair as alt

# Optional DOCX dependency (preferred over PDF for reliability)
try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ---------- Column mapping tailored to your Excel ----------
COLUMN_MAP = {
    "Code": "code",
    "Title": "title",
    "Opening Date": "opening_date",
    "Deadline": "deadline",
    "First Stage Deadline": "first_deadline",
    "Second Stage Deadline": "second_deadline",
    "Two-Stage": "two_stage",
    "Cluster": "cluster",
    "Destination": "destination_or_strand",
    "Budget Per Project": "budget_per_project_eur",
    "Total Budget": "total_budget_eur",
    "Number of Projects": "num_projects",
    "Type of Action": "type_of_action",
    "TRL": "trl",
    "Call Name": "call_name",
    "Expected Outcome": "expected_outcome",
    "Scope": "scope",
    "Description": "full_text",
    "Source Filename": "source_filename",
    "Version Label": "version_label",
    "Parsed On (UTC)": "parsed_on_utc",
}

DISPLAY_COLS = [
    "code","title","opening_date",
    "deadline","first_deadline","second_deadline","two_stage",
    "cluster","destination_or_strand","type_of_action","trl",
    "budget_per_project_eur","total_budget_eur","num_projects",
    "call_name","version_label","source_filename","parsed_on_utc",
]

SEARCHABLE_COLUMNS = (
    "code","title","call_name","expected_outcome","scope","full_text",
    "cluster","destination_or_strand","type_of_action","trl"
)

# ---------- Helpers ----------
def nl_to_br(s: str) -> str:
    return "" if not s else s.replace("\n", "<br>")

def clean_footer(text: str) -> str:
    """
    Remove footer lines like:
    '... Work Programme 2026-2027 ... Page xx of yy ...'
    even if embedded in a longer line.
    """
    if not text:
        return ""
    footer_pattern = re.compile(
        r"Horizon\s*Europe\s*[-–]?\s*Work Programme.*?Page\s+\d+\s+of\s+\d+",
        re.IGNORECASE | re.DOTALL
    )
    cleaned = footer_pattern.sub("", text)
    return re.sub(r"\s+", " ", cleaned).strip()

def normalize_bullets(text: str) -> str:
    if not isinstance(text, str) or text == "":
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?m)^[ \t]*[▪◦●•]\s*", "- ", text)  # only at line start
    text = re.sub(r"[ \t]+", " ", text)                  # collapse spaces, keep \n
    text = re.sub(r"(?<!\n)([ \t]+[-*]\s+)", r"\n- ", text)
    text = re.sub(r"(?<!\n)([ \t]+)(\d+\.\s+)", r"\n\2", text)
    return text.strip()

def highlight_text(text: str, keywords: list[str], colours=None) -> str:
    """Return text with keywords highlighted using HTML span tags."""
    if not text:
        return ""
    clean_keywords = [str(k).strip() for k in keywords if k and str(k).strip()]
    if not clean_keywords:
        return text
    if colours is None:
        colours = ["#ffff00", "#a0e7e5", "#ffb3b3"]  # yellow, teal, pink
    highlighted = str(text)
    for i, kw in enumerate(clean_keywords):
        colour = colours[i % len(colours)]
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        highlighted = pattern.sub(
            lambda m: f"<span style='background-color:{colour}; font-weight:bold;'>{m.group(0)}</span>",
            highlighted
        )
    return highlighted

def canonicalise(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = [c.strip() for c in df.columns]
    for src, dst in COLUMN_MAP.items():
        if src in df.columns and dst not in df.columns:
            df = df.rename(columns={src: dst})
    df = df.rename(columns={c: c.strip().lower() for c in df.columns})
    if "programme" not in df.columns:
        df["programme"] = "Horizon Europe"
    for c in ("opening_date","deadline","first_deadline","second_deadline"):
        if c in df.columns:
            df[c] = pd.to_datetime(df[c], errors="coerce", dayfirst=True)
    for c in ("budget_per_project_eur","total_budget_eur","trl","num_projects"):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")
    if "two_stage" in df.columns:
        df["two_stage"] = (
            df["two_stage"].astype(str).str.strip().str.lower()
            .map({"true": True, "false": False, "yes": True, "no": False})
            .fillna(False)
        )
    else:
        df["two_stage"] = False
    return df

def wrap_label(text: str, width=50, max_lines=3) -> str:
    s = str(text or "")
    parts = [s[i:i+width] for i in range(0, len(s), width)]
    return "\n".join(parts[:max_lines])

def safe_date_bounds(series, start_fb="2000-01-01", end_fb="2100-12-31"):
    s = pd.to_datetime(series, errors="coerce").dropna()
    if s.empty:
        return (pd.to_datetime(start_fb).date(), pd.to_datetime(end_fb).date())
    lo, hi = s.min().date(), s.max().date()
    if lo == hi:
        hi = (pd.to_datetime(hi) + pd.Timedelta(days=1)).date()
    return lo, hi

def build_month_bands(min_x: pd.Timestamp, max_x: pd.Timestamp) -> pd.DataFrame:
    start = pd.Timestamp(min_x).to_period("M").start_time
    end   = (pd.Timestamp(max_x).to_period("M") + 1).start_time
    months = pd.date_range(start, end, freq="MS")
    rows = []
    for i in range(len(months) - 1):
        rows.append({"start": months[i], "end": months[i+1], "band": i % 2})
    return pd.DataFrame(rows)

# -------- Vectorised multi-keyword search --------
def multi_keyword_filter(df: pd.DataFrame, terms: list[str], mode: str, title_code_only: bool) -> pd.DataFrame:
    terms = [t.strip().lower() for t in terms if t and str(t).strip()]
    if not terms:
        return df
    hay = df["_search_title"] if title_code_only else df["_search_all"]
    if mode.upper() == "AND":
        pattern = "".join([f"(?=.*{re.escape(t)})" for t in terms]) + ".*"
    else:
        pattern = "(" + "|".join(re.escape(t) for t in terms) + ")"
    mask = hay.str.contains(pattern, regex=True, na=False)
    return df[mask]

# -------- Build long-form segments for Altair (two-stage on SAME ROW) --------
def build_segments(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for _, r in df.iterrows():
        code = str(r.get("code") or "")
        title = str(r.get("title") or "")
        y_label = wrap_label(f"{code}", width=100, max_lines=5)

        prog = r.get("programme")
        clu  = r.get("cluster")
        open_dt   = r.get("opening_date")
        final_dt  = r.get("deadline")
        first_dt  = r.get("first_deadline")
        second_dt = r.get("second_deadline")
        two_stage = bool(r.get("two_stage"))
        title_inbar = wrap_label(title, width=100, max_lines=3)

        if two_stage:
            if pd.notna(open_dt) and pd.notna(first_dt) and open_dt <= first_dt:
                bar_days = (first_dt - open_dt).days
                rows.append({
                    "y_label": y_label, "programme": prog,
                    "start": open_dt, "end": first_dt,
                    "segment": "Stage 1",
                    "title": title, "title_inbar": title_inbar,
                    "budget_per_project_eur": r.get("budget_per_project_eur"),
                    "type_of_action" : r.get("type_of_action"),
                    "cluster": clu,
                    "bar_days": bar_days,
                    "mid": open_dt + (first_dt - open_dt)/2,
                })
            segB_end = second_dt if pd.notna(second_dt) else (final_dt if pd.notna(final_dt) else None)
            if pd.notna(first_dt) and pd.notna(segB_end) and first_dt <= segB_end:
                bar_days = (segB_end - first_dt).days
                rows.append({
                    "y_label": y_label, "programme": prog,
                    "start": first_dt, "end": segB_end,
                    "segment": "Stage 2",
                    "title": title,"title_inbar": "",
                    "budget_per_project_eur": r.get("budget_per_project_eur"),
                    "type_of_action" : r.get("type_of_action"),
                    "cluster": clu,
                    "bar_days": bar_days,
                    "mid": first_dt + (segB_end - first_dt)/2,
                })
        else:
            if pd.notna(open_dt) and pd.notna(final_dt) and open_dt <= final_dt:
                bar_days = (final_dt - open_dt).days
                rows.append({
                    "y_label": y_label, "programme": prog,
                    "start": open_dt, "end": final_dt,
                    "segment": "Single",
                    "title": title, "title_inbar": title_inbar,
                    "budget_per_project_eur": r.get("budget_per_project_eur"),
                    "type_of_action" : r.get("type_of_action"),
                    "cluster": clu,
                    "bar_days": bar_days,
                    "mid": open_dt + (final_dt - open_dt)/2,
                })

    seg = pd.DataFrame(rows)
    if seg.empty:
        return seg
    seg["earliest_end"] = seg.groupby("y_label")["end"].transform("min")
    seg = seg.sort_values(["earliest_end", "start"]).reset_index(drop=True)
    return seg

def build_altair_chart_from_segments(seg: pd.DataFrame, view_start, view_end):
    if seg.empty:
        return None
    y_order = seg["y_label"].drop_duplicates().tolist()
    unique_rows = len(y_order)
    row_height = 46
    bar_size     = int(row_height * 0.38)               # bar thickness ~38% of band
    label_offset = - int(bar_size / 2 + 4)              # position text just ABOVE the bar

    chart_height = max(1500, unique_rows * row_height)
    domain_min = seg["start"].min()
    domain_max = seg["end"].max()
    min_x = min(seg["start"].min(), seg["end"].min())
    max_x = max(seg["start"].max(), seg["end"].max())
    bands_df = build_month_bands(min_x, max_x)

    month_shade = (
        alt.Chart(bands_df)
        .mark_rect(tooltip=False)
        .encode(
            x=alt.X("start:T"),
            x2=alt.X2("end:T"),
            opacity=alt.Opacity("band:Q",
                                scale=alt.Scale(domain=[0,1], range=[0.0, 0.15]),
                                legend=None),
            color=alt.value("#808080")
        )
    )

    months = pd.date_range(pd.Timestamp(min_x).to_period("M").start_time,
                           pd.Timestamp(max_x).to_period("M").end_time,
                           freq="MS")
    month_grid = alt.Chart(pd.DataFrame({"t": months})).mark_rule(stroke="#000000", strokeWidth=0.5).encode(x="t:T")

    month_labels_df = pd.DataFrame({
        "month": months[:-1],
        "next_month": months[1:],
        "label": [m.strftime("%b %Y") for m in months[:-1]]
    })
    month_labels_df["mid"] = month_labels_df["month"] + (
        (month_labels_df["next_month"] - month_labels_df["month"]) / 2
    )
    month_labels = alt.Chart(month_labels_df).mark_text(
        align="center", baseline="top", dy=0, fontSize=12, fontWeight="bold",
    ).encode(x="mid:T", text="label:N", y=alt.value(0))

    # Thin top axis rule to visually reinforce the top axis
    # top_axis_rule = alt.Chart(pd.DataFrame({"t":[domain_min, domain_max]})).mark_rule(stroke="#333", strokeWidth=1).encode(
        #x="t:T"
    #)

    base = alt.Chart(seg).encode(
        y=alt.Y(
            "y_label:N",
            sort=y_order,
            axis=alt.Axis(
                title=None, labelLimit=200, labelFontSize=11, labelAlign="right",
                labelPadding=50, domain=True
            ),
            scale=alt.Scale(domain=y_order, paddingInner=0.3, paddingOuter=0.8)
        )
    )

    # Bars: colour ONLY via encoding so Stage-2 opacity works (no fixed mark color)
    bars = alt.Chart(seg).mark_bar(cornerRadius=10, size=bar_size).encode(
        y=alt.Y(
            "y_label:N",
            sort=y_order,
            axis=alt.Axis(
                title=None,
                labelLimit=500,
                labelFontSize=13,
                labelAlign="right",
                labelPadding=50,
                domain=True
            ),
            scale=alt.Scale(domain=y_order, paddingInner=0.3, paddingOuter=0.8)
        ),
        x=alt.X(
            "start:T",
            axis=alt.Axis(
                title=None, format="%b %Y", tickCount="month", orient="top",
                labelFontSize=11, labelPadding=50, labelOverlap="greedy", tickSize=6
            ),
            scale=alt.Scale(domain=[domain_min, domain_max])
        ),
        x2="end:T",
        color=alt.Color(
            "type_of_action:N",
            legend=alt.Legend(title="Type of Action",
                              orient="top",
                              direction="horizontal",
                              offset=100),
            scale=alt.Scale(scheme="set2")
        ),
        opacity=alt.condition(alt.datum.segment == "Stage 2", alt.value(0.7), alt.value(1.0)),
        tooltip=[
            alt.Tooltip("title:N", title="Title"),
            alt.Tooltip("programme:N", title="Programme"),
            alt.Tooltip("budget_per_project_eur:Q", title="Budget (€)", format=",.0f"),
            alt.Tooltip("start:T", title="Start", format="%d %b %Y"),
            alt.Tooltip("end:T", title="End", format="%d %b %Y"),
            alt.Tooltip("segment:N", title="Segment")
        ]
    )

    start_labels = base.mark_text(align="right",
                                  dx=-4,
                                  dy=5,
                                  fontSize=10,
                                  color="#111")\
        .encode(x="start:T",
                text=alt.Text("start:T",
                              format="%d %b %Y"))
    end_labels   = base.mark_text(align="left",
                                  dx=4,
                                  dy=5,
                                  fontSize=10,
                                  color="#111")\
        .encode(x="end:T",
                text=alt.Text("end:T",
                              format="%d %b %Y"))

    text_cond = alt.condition(alt.datum.bar_days >= 10,
                              alt.value(1),
                              alt.value(0))
    
    inbar = base.mark_text(
        align="left",
        baseline="bottom",
        dx=2,
        dy=label_offset,
        color="black"   # styling is fine here
    ).encode(
        x=alt.X("start:T", scale=alt.Scale(domain=[domain_min, domain_max]), axis=None),
        text=alt.Text("title_inbar:N"),
        opacity=alt.condition(alt.datum.bar_days >= 10, alt.value(1), alt.value(0))
    )

    # --- Today line (Europe/Brussels), drawn as a dashed red rule with a tooltip
    today_ts = pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None)
    today_df = pd.DataFrame({"t": [today_ts]})
    
    today_rule = alt.Chart(today_df).mark_rule(
        color="#d62728", strokeDash=[6,4], strokeWidth=2
    ).encode(
        x="t:T",
        tooltip=[alt.Tooltip("t:T", title="Today", format="%d %b %Y")]
    )
    
    today_label = alt.Chart(today_df).mark_text(
        align="left", baseline="top", dx=4, dy=0, fontSize=11, fontWeight="bold", color="#d62728"
    ).encode(
        x="t:T",
        y=alt.value(0),     # inside top edge
        text=alt.Text("t:T", format= 'Today: "%d %b %Y"')
    )
    
    chart = (
        month_shade + month_grid + bars + start_labels + end_labels + inbar + month_labels + today_rule + today_label
    ).properties(
        height= max(800, unique_rows * row_height),
        width='container',  # fill available width
        padding={"top": 50, "bottom": 30, "left": 10, "right": 10}
    ).configure_axis(
        grid=False,
        domain=True,
        domainWidth=1
    ).configure_view(
        continuousHeight=500, continuousWidth=500, strokeWidth=0, clip=False,
    ).resolve_scale(
        x='shared', y='shared'
    ).resolve_axis(
        x='shared', y='shared'
    ).interactive(bind_x=True)  # horizontal pan/zoom

    return chart

# ---------- Cached I/O & options ----------
@st.cache_data(show_spinner=False)
def get_sheet_names(file_bytes: bytes) -> List[str]:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    return xls.sheet_names

@st.cache_data(show_spinner=False)
def load_sheet(file_bytes: bytes, sheet_name: str) -> pd.DataFrame:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    raw = pd.read_excel(xls, sheet_name=sheet_name)
    df = canonicalise(raw)

    # Build searchable haystacks once
    present = [c for c in SEARCHABLE_COLUMNS if c in df.columns]
    if present:
        df["_search_all"] = df[present].astype(str).agg(" ".join, axis=1).str.lower()
    else:
        df["_search_all"] = ""
    title_cols = [c for c in ["code", "title"] if c in df.columns]
    if title_cols:
        df["_search_title"] = df[title_cols].astype(str).agg(" ".join, axis=1).str.lower()
    else:
        df["_search_title"] = ""

    # Convenience "any closing" column for filtering and sorting
    close_cols = [c for c in ["deadline","first_deadline","second_deadline"] if c in df.columns]
    if close_cols:
        df["closing_date_any"] = pd.to_datetime(
            df[close_cols].stack(), errors="coerce"
        ).groupby(level=0).min()
    else:
        df["closing_date_any"] = pd.NaT

    return df

@st.cache_data(show_spinner=False)
def derive_filter_options(df: pd.DataFrame):
    prog_opts    = sorted([p for p in df["programme"].dropna().unique().tolist() if p != ""])
    cluster_opts = sorted([c for c in df.get("cluster", pd.Series(dtype=object)).dropna().unique().tolist() if c != ""])
    type_opts    = sorted([t for t in df.get("type_of_action", pd.Series(dtype=object)).dropna().unique().tolist() if t != ""])
    trl_opts     = sorted([str(int(x)) for x in df.get("trl", pd.Series(dtype=float)).dropna().unique() if pd.notna(x)])
    dest_opts    = sorted([d for d in df.get("destination_or_strand", pd.Series(dtype=object)).dropna().unique().tolist() if d != ""])
    return prog_opts, cluster_opts, type_opts, trl_opts, dest_opts

# ---------- DOCX / HTML report ----------
def generate_docx_report(calls_df: pd.DataFrame, notes_by_code: Dict[str, str], title: str="Funding Report") -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed")

    doc = Document()
    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph(f"Generated on {datetime.utcnow():%d %b %Y, %H:%M UTC}")
    p.runs[0].font.size = Pt(9)

    # Summary table
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    for i, t in enumerate(["Code", "Title", "Opening", "Deadline"]):
        hdr[i].text = t

    for _, r in calls_df.iterrows():
        opening = r.get("opening_date"); deadline = r.get("deadline")
        opening_s = opening.strftime("%d %b %Y") if pd.notna(opening) else "-"
        deadline_s = deadline.strftime("%d %b %Y") if pd.notna(deadline) else "-"
        row = table.add_row().cells
        row[0].text = str(r.get("code",""))
        row[1].text = str(r.get("title",""))
        row[2].text = opening_s
        row[3].text = deadline_s

    # Detailed sections
    for _, r in calls_df.iterrows():
        doc.add_page_break()
        doc.add_heading(f"{r.get('code','')} — {r.get('title','')}", level=1)

        lines = []
        lines.append(f"Programme: {r.get('programme','-')}")
        lines.append(f"Cluster: {r.get('cluster','-')}")
        lines.append(f"Destination: {r.get('destination_or_strand','-')}")
        lines.append(f"Type of Action: {r.get('type_of_action','-')}")
        trl_val = r.get("trl")
        lines.append(f"TRL: {int(trl_val) if pd.notna(trl_val) else '-'}")
        opening = r.get("opening_date"); deadline = r.get("deadline")
        first_deadline = r.get("first_deadline"); second_deadline = r.get("second_deadline")
        lines.append(f"Opening: {opening:%d %b %Y}" if pd.notna(opening) else "Opening: -")
        lines.append(f"Deadline: {deadline:%d %b %Y}" if pd.notna(deadline) else "Deadline: -")
        if r.get("two_stage"):
            lines.append(f"Stage 1: {first_deadline:%d %b %Y}" if pd.notna(first_deadline) else "Stage 1: -")
            lines.append(f"Stage 2: {second_deadline:%d %b %Y}" if pd.notna(second_deadline) else "Stage 2: -")
        bpp = r.get("budget_per_project_eur"); tot = r.get("total_budget_eur"); npj = r.get("num_projects")
        lines.append(f"Budget per project: {bpp:,.0f} EUR" if pd.notna(bpp) else "Budget per project: -")
        lines.append(f"Total budget: {tot:,.0f} EUR" if pd.notna(tot) else "Total budget: -")
        lines.append(f"# Projects: {int(npj) if pd.notna(npj) else '-'}")

        doc.add_paragraph("\n".join(lines))

        notes = (notes_by_code or {}).get(str(r.get("code","")), "")
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(notes if notes else "-")

        # Optional long text (if provided in df)
        eo = clean_footer(str(r.get("expected_outcome") or ""))
        sc = clean_footer(str(r.get("scope") or ""))
        if eo:
            doc.add_heading("Expected Outcome", level=2)
            for line in normalize_bullets(eo).splitlines():
                if line.startswith("- "):
                    par = doc.add_paragraph(line[2:])
                    par.paragraph_format.left_indent = Cm(0.5)
                else:
                    doc.add_paragraph(line)
        if sc:
            doc.add_heading("Scope", level=2)
            for line in normalize_bullets(sc).splitlines():
                if line.startswith("- "):
                    par = doc.add_paragraph(line[2:])
                    par.paragraph_format.left_indent = Cm(0.5)
                else:
                    doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def generate_html_report(calls_df: pd.DataFrame, notes_by_code: Dict[str, str], title: str="Funding Report") -> bytes:
    """Fallback lightweight HTML report that users can save as PDF from the browser."""
    parts = [f"<h1>{title}</h1><p><em>Generated on {datetime.utcnow():%d %b %Y, %H:%M UTC}</em></p>"]
    # Summary table
    parts.append("<table border='1' cellspacing='0' cellpadding='4'><tr><th>Code</th><th>Title</th><th>Opening</th><th>Deadline</th></tr>")
    for _, r in calls_df.iterrows():
        opening = r.get("opening_date"); deadline = r.get("deadline")
        opening_s = opening.strftime("%d %b %Y") if pd.notna(opening) else "-"
        deadline_s = deadline.strftime("%d %b %Y") if pd.notna(deadline) else "-"
        parts.append(f"<tr><td>{r.get('code','')}</td><td>{r.get('title','')}</td><td>{opening_s}</td><td>{deadline_s}</td></tr>")
    parts.append("</table>")

    # Detailed sections
    for _, r in calls_df.iterrows():
        parts.append("<hr>")
        parts.append(f"<h2>{r.get('code','')} — {r.get('title','')}</h2>")
        meta = []
        meta.append(f"Programme: {r.get('programme','-')}")
        meta.append(f"Cluster: {r.get('cluster','-')}")
        meta.append(f"Destination: {r.get('destination_or_strand','-')}")
        meta.append(f"Type of Action: {r.get('type_of_action','-')}")
        trl_val = r.get("trl")
        meta.append(f"TRL: {int(trl_val) if pd.notna(trl_val) else '-'}")
        opening = r.get("opening_date"); deadline = r.get("deadline")
        first_deadline = r.get("first_deadline"); second_deadline = r.get("second_deadline")
        meta.append(f"Opening: {opening:%d %b %Y}" if pd.notna(opening) else "Opening: -")
        meta.append(f"Deadline: {deadline:%d %b %Y}" if pd.notna(deadline) else "Deadline: -")
        if r.get("two_stage"):
            meta.append(f"Stage 1: {first_deadline:%d %b %Y}" if pd.notna(first_deadline) else "Stage 1: -")
            meta.append(f"Stage 2: {second_deadline:%d %b %Y}" if pd.notna(second_deadline) else "Stage 2: -")
        bpp = r.get("budget_per_project_eur"); tot = r.get("total_budget_eur"); npj = r.get("num_projects")
        meta.append(f"Budget per project: {bpp:,.0f} EUR" if pd.notna(bpp) else "Budget per project: -")
        meta.append(f"Total budget: {tot:,.0f} EUR" if pd.notna(tot) else "Total budget: -")
        meta.append(f"# Projects: {int(npj) if pd.notna(npj) else '-'}")
        parts.append("<p>" + "<br>".join(meta) + "</p>")

        notes = (notes_by_code or {}).get(str(r.get("code","")), "")
        parts.append("<p><b>Notes</b><br>" + (notes.replace("\n", "<br>") if notes else "-") + "</p>")

        eo = clean_footer(str(r.get("expected_outcome") or ""))
        sc = clean_footer(str(r.get("scope") or ""))
        if eo:
            parts.append("<p><b>Expected Outcome</b><br>" + normalize_bullets(eo).replace("\n","<br>") + "</p>")
        if sc:
            parts.append("<p><b>Scope</b><br>" + normalize_bullets(sc).replace("\n","<br>") + "</p>")

    html = "<html><head><meta charset='utf-8'><title>{}</title></head><body>{}</body></html>".format(title, "".join(parts))
    return html.encode("utf-8")

# ---------- UI ----------
st.set_page_config(page_title="Funding Dashboard", layout="wide")

# Global CSS: widen page & provide scrolling container
st.markdown(
    """
    <style>
    .scroll-container {
        overflow-x: auto;
        overflow-y: auto;
        max-height: 900px;          /* vertical scroll for tall charts */
        padding: 16px;
        border: 1px solid #eee;
        border-radius: 8px;
    }
    .main .block-container { padding-left: 1.5rem; padding-right: 1.5rem; max-width: 95vw; }
    </style>
    """,
    unsafe_allow_html=True
)

# Logo (optional)
try:
    with open("logo.png", "rb") as f:
        data_b64 = base64.b64encode(f.read()).decode("utf-8")
    st.markdown(
        f"""
        <div style="text-align: center;">
            <img src="data:image/png;base64,{data_b64}" width="250">
        </div>
        """,
        unsafe_allow_html=True
    )
except Exception:
    pass

st.title("Funding Dashboard")

st.info(
    "📂 Please upload the latest parsed Excel file.\n\n"
    "➡️ Location hint:\n\n"
    "- **3.SA Practices** → Central Systems and Bid Management → 1. Central Systems → 2. CS EU PROGRAMMES Area → 4. Horizon Europe → CALENDAR OF CALLS FOR PROPOSALS → 4. WPs 2026-2027 → Coding Davide\n\n"
    "👉 Look for *Horizon Calls - 26_27*.\n"
)

upl = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])
if not upl:
    st.stop()

# Sheet selection
sheet_names = get_sheet_names(upl.getvalue())
sheet = st.selectbox("Sheet", sheet_names, index=0)

# Load data
df = load_sheet(upl.getvalue(), sheet)

# Compute date bounds ONCE
open_lo, open_hi = safe_date_bounds(df.get("opening_date"))
dead_all = pd.concat([
    pd.to_datetime(df.get("deadline"), errors="coerce"),
    pd.to_datetime(df.get("first_deadline"), errors="coerce"),
    pd.to_datetime(df.get("second_deadline"), errors="coerce"),
], axis=0)
dead_lo, dead_hi = safe_date_bounds(dead_all)

# ----- Top filter form -----
prog_opts, cluster_opts, type_opts, trl_opts, dest_opts = derive_filter_options(df)

with st.form("filters_form", clear_on_submit=False):
    st.header("Filters")

    col1, col2, col3 = st.columns(3)
    with col1:
        programmes = st.multiselect("Programme", options=prog_opts, default=prog_opts)
    with col2:
        clusters   = st.multiselect("Cluster", options=cluster_opts)
    with col3:
        dests      = st.multiselect("Destination / Strand", options=dest_opts)

    col4, col5 = st.columns(2)
    with col4:
        types      = st.multiselect("Type of Action", options=type_opts)
    with col5:
        trls       = st.multiselect("TRL", options=trl_opts)

    # Keyword row
    col6, col7, col8, col9 = st.columns([2,2,2,1])
    with col6:
        kw1 = st.text_input("Keyword 1")
    with col7:
        kw2 = st.text_input("Keyword 2")
    with col8:
        kw3 = st.text_input("Keyword 3")
    with col9:
        combine_mode = st.radio("Combine", ["AND", "OR"], horizontal=True, index=0)
    title_code_only = st.checkbox("Search only in Title & Code", value=True)

    # Date filters row
    col10, col11, col12, col13 = st.columns(4)
    with col10:
        open_start = st.date_input("Open from", value=open_lo, min_value=open_lo, max_value=open_hi)
    with col11:
        open_end   = st.date_input("Open to",   value=open_hi, min_value=open_lo, max_value=open_hi)
    with col12:
        close_from = st.date_input("Close from", value=dead_lo, min_value=dead_lo, max_value=dead_hi)
    with col13:
        close_to   = st.date_input("Close to",   value=dead_hi, min_value=dead_lo, max_value=dead_hi)

    # Budget slider (dynamic step)
    bud_series = pd.to_numeric(df.get("budget_per_project_eur"), errors="coerce").dropna()
    if bud_series.empty:
        min_bud, max_bud = 0.0, 1_000_000.0
    else:
        min_bud, max_bud = float(bud_series.min()), float(bud_series.max())
        if not (min_bud < max_bud):
            min_bud, max_bud = max(min_bud, 0.0), min_bud + 100000.0
    rng = max_bud - min_bud
    try:
        step = max(1e4, round(rng / 50, -3))  # ~50 steps, nearest 1k
    except Exception:
        step = 10000.0
    budget_range = st.slider("Budget per project (EUR)", min_bud, max_bud, (min_bud, max_bud), step=step)

    applied = st.form_submit_button("Apply filters")

# Persist criteria
if "criteria" not in st.session_state:
    st.session_state.criteria = {}
if applied:
    st.session_state.criteria = dict(
        programmes=programmes, clusters=clusters, types=types, trls=trls, dests=dests,
        kw1=kw1, kw2=kw2, kw3=kw3, combine_mode=combine_mode, title_code_only=title_code_only,
        open_start=open_start, open_end=open_end, close_from=close_from, close_to=close_to,
        budget_range=budget_range, applied_at=datetime.utcnow().strftime("%H:%M UTC")
    )

if not st.session_state.criteria:
    st.session_state.criteria = dict(
        programmes=sorted(df["programme"].dropna().unique().tolist()),
        clusters=[], types=[], trls=[], dests=[],
        kw1="", kw2="", kw3="", combine_mode="OR", title_code_only=False,
        open_start=open_lo, open_end=open_hi, close_from=dead_lo, close_to=dead_hi,
        budget_range=(0.0, 1_000_000.0), applied_at=None
    )

crit = st.session_state.criteria
if crit.get("applied_at"):
    st.caption(f"Filters last applied at {crit['applied_at']}")

# ---- Apply filters ----
f = df.copy()

# Keywords
f = multi_keyword_filter(f, [crit["kw1"], crit["kw2"], crit["kw3"]], crit["combine_mode"], crit["title_code_only"])

# Categorical filters
if crit["programmes"]: f = f[f["programme"].isin(crit["programmes"])]
if crit["clusters"]:   f = f[f["cluster"].isin(crit["clusters"])]
if crit["types"]:      f = f[f["type_of_action"].isin(crit["types"])]
if crit["trls"]:
    f = f[f["trl"].dropna().astype("Int64").astype(str).isin(crit["trls"])]
if crit["dests"]:      f = f[f["destination_or_strand"].isin(crit["dests"])]

# Dates
f = f[f["opening_date"].between(pd.to_datetime(crit["open_start"]), pd.to_datetime(crit["open_end"]), inclusive="both")]
f = f[f["closing_date_any"].between(pd.to_datetime(crit["close_from"]), pd.to_datetime(crit["close_to"]), inclusive="both")]

# Budget
low, high = crit["budget_range"]
f = f[f["budget_per_project_eur"].fillna(0).between(low, high)]

st.markdown(f"**Showing {len(f)} rows** after last applied filters.")

# ---------- Tabs ----------
tab1, tab2, tab3, tab4 = st.tabs(["📅 Gantt", "📋 Table", "📚 Full Data", "📝 Shortlist & Notes (DOCX)"])

with tab1:
    st.subheader("Gantt (Opening → Stage 1 → Stage 2 / Final)")
    segments = build_segments(f)
    if segments.empty:
        st.info("No rows with valid dates to display.")
    else:
        # Controls
        group_mode = st.radio("Group charts by", ["None", "Cluster"], horizontal=True, index=0)
        view_mode = st.radio(
            "View",
            ["Dropdowns (one per group)", "Single select (one chart)"],
            horizontal=True, index=0,
            help="Dropdowns show all groups as expanders; Single select renders only one chart."
        )

        def render_chart(seg_df, title_suffix=""):
            chart = build_altair_chart_from_segments(
                seg_df,
                view_start=crit["open_start"],
                view_end=crit["close_to"]
            )
            if title_suffix:
                st.markdown(f"### {title_suffix}")
            st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
            st.altair_chart(chart, use_container_width=True)  # fill width; scroll container handles overflow
            st.markdown('</div>', unsafe_allow_html=True)

        if group_mode == "None":
            render_chart(segments)
        else:
            key = "cluster"
            if key not in segments.columns:
                st.warning(f"Column '{key}' not available in data.")
            else:
                grouped = list(segments.groupby(key))
                grouped.sort(key=lambda kv: len(kv[1]), reverse=True)

                if view_mode.startswith("Single"):
                    names = [str(k if pd.notna(k) else "—") for k, _ in grouped]
                    sel = st.selectbox("Select cluster", options=names, index=0)
                    for (name, gdf), disp in zip(grouped, names):
                        if disp == sel:
                            render_chart(gdf, f"Cluster: {disp} ({len(gdf)} calls)")
                            break
                else:
                    for name, gdf in grouped:
                        disp = str(name if pd.notna(name) else "—")
                        with st.expander(f"Cluster: {disp} ({len(gdf)} calls)", expanded=False):
                            render_chart(gdf)

with tab2:
    st.subheader("Filtered table")
    show_cols = [c for c in DISPLAY_COLS if c in f.columns]
    group_by_cluster = st.checkbox("Group by Cluster")
    if group_by_cluster and "cluster" in f.columns:
        for clu, group_df in f.groupby("cluster"):
            with st.expander(f"Cluster: {clu} ({len(group_df)} calls)"):
                st.dataframe(group_df[show_cols], use_container_width=True, hide_index=True, height=400)
    else:
        st.dataframe(f[show_cols], use_container_width=True, hide_index=True, height=800)

    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as xw:
        f.to_excel(xw, index=False, sheet_name="filtered")
    out.seek(0)
    src_label = f.get("version_label").dropna().astype(str).unique() if "version_label" in f.columns else []
    suffix = src_label[0] if isinstance(src_label, (list, pd.Series)) and len(src_label)==1 else "filtered"
    st.download_button("⬇️ Download filtered (Excel)", out,
                       file_name=f"calls_{suffix}.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

with tab3:
    st.subheader("Full data (expand rows)")
    group_full_by_cluster = st.checkbox("Group by Cluster (full data)")

    kw_list = [crit.get("kw1", ""), crit.get("kw2", ""), crit.get("kw3", "")]

    def render_row(row):
        col1, col2 = st.columns(2)
        with col1:
            if pd.notna(row.get("opening_date")):
                st.markdown(f"📅 **Opening:** {row.get('opening_date'):%d %b %Y}")
            if pd.notna(row.get("deadline")):
                st.markdown(f"⏳ **Deadline:** {row.get('deadline'):%d %b %Y}")
            if row.get("two_stage"):
                if pd.notna(row.get("first_deadline")):
                    st.markdown(f"🔄 **Stage 1:** {row.get('first_deadline'):%d %b %Y}")
                if pd.notna(row.get("second_deadline")):
                    st.markdown(f"🔄 **Stage 2:** {row.get('second_deadline'):%d %b %Y}")
        with col2:
            if pd.notna(row.get("budget_per_project_eur")):
                st.markdown(f"💶 **Budget per project:** {row.get('budget_per_project_eur'):,.0f} EUR")
            if pd.notna(row.get("total_budget_eur")):
                st.markdown(f"📦 **Total budget:** {row.get('total_budget_eur'):,.0f} EUR")
            if pd.notna(row.get("num_projects")):
                st.markdown(f"📊 **# Projects:** {int(row.get('num_projects'))}")

        st.markdown(
            f"🏷️ **Programme:** {row.get('programme','-')}  "
            f"| **Cluster:** {row.get('cluster','-')}  "
            f"| **Destination:** {row.get('destination_or_strand','-')}  "
            f"| **Type of Action:** {row.get('type_of_action','-')}  "
            f"| **TRL:** {row.get('trl','-')}"
        )

        if row.get("expected_outcome"):
            with st.expander("🎯 Expected Outcome"):
                clean_text = normalize_bullets(clean_footer(row.get("expected_outcome")))
                clean_text = nl_to_br(clean_text)
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)

        if row.get("scope"):
            with st.expander("🧭 Scope"):
                clean_text = normalize_bullets(clean_footer(row.get("scope")))
                clean_text = nl_to_br(clean_text)
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)

        if row.get("full_text"):
            with st.expander("📖 Full Description"):
                clean_text = normalize_bullets(clean_footer(row.get("full_text")))
                clean_text = nl_to_br(clean_text)
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)

        st.caption(
            f"📂 Source: {row.get('source_filename','-')} "
            f"| Version: {row.get('version_label','-')} "
            f"| Parsed on: {row.get('parsed_on_utc','-')}"
        )

    if group_full_by_cluster and "cluster" in f.columns:
        for clu, group_df in f.groupby("cluster"):
            with st.expander(f"Cluster: {clu} ({len(group_df)} calls)"):
                for _, row in group_df.iterrows():
                    title = f"{row.get('code','')} — {row.get('title','')}"
                    with st.expander(title):
                        render_row(row)
    else:
        for _, row in f.iterrows():
            title = f"{row.get('code','')} — {row.get('title','')}"
            with st.expander(title):
                render_row(row)

with tab4:
    st.subheader("Shortlist & Notes → Generate Report (DOCX/HTML)")

    # 1) Show list of codes/titles based on current filters
    if "selection" not in st.session_state:
        st.session_state.selection = set()
    if "notes" not in st.session_state:
        st.session_state.notes = {}

    st.markdown("**Select calls to include in the report**")
    for _, r in f.sort_values(["closing_date_any", "opening_date"]).iterrows():
        code = str(r.get("code") or "")
        title = str(r.get("title") or "")
        label = f"{code} — {title}"
        checked = code in st.session_state.selection
        new_val = st.checkbox(label, value=checked, key=f"chk_{code}")
        if new_val and not checked:
            st.session_state.selection.add(code)
        elif (not new_val) and checked:
            st.session_state.selection.discard(code)

    # 2) Notes per selected call
    if st.session_state.selection:
        st.markdown("---")
        st.markdown("**Enter notes per selected call**")
        selected_df = f[f["code"].astype(str).isin(st.session_state.selection)].copy()

        for _, r in selected_df.iterrows():
            code = str(r.get("code") or "")
            title = str(r.get("title") or "")
            default_txt = st.session_state.notes.get(code, "")
            st.session_state.notes[code] = st.text_area(
                f"Notes — {code} — {title}",
                value=default_txt, key=f"note_{code}", height=120
            )

        st.markdown("---")
        colA, colB = st.columns(2)
        with colA:
            report_title = st.text_input("Report title", value="Funding Report – Shortlist")
        with colB:
            include_long_text = st.checkbox("Include Expected Outcome / Scope (export)", value=False,
                                            help="If off, those sections are omitted for a shorter document.")

        # 3) Generate report (DOCX preferred, else HTML)
        def prep_df_for_report(df_in: pd.DataFrame) -> pd.DataFrame:
            cols = [
                "code","title","programme","cluster","destination_or_strand",
                "type_of_action","trl","opening_date","deadline",
                "first_deadline","second_deadline","two_stage",
                "budget_per_project_eur","total_budget_eur","num_projects",
                "expected_outcome","scope"
            ]
            keep = [c for c in cols if c in df_in.columns]
            return df_in[keep].copy()

        report_df = prep_df_for_report(selected_df)

        if st.button("📄 Generate report"):
            try:
                df_for_export = report_df.copy()
                if not include_long_text:
                    for col in ("expected_outcome","scope"):
                        if col in df_for_export.columns:
                            df_for_export[col] = ""

                if DOCX_AVAILABLE:
                    docx_bytes = generate_docx_report(df_for_export, st.session_state.notes, title=report_title)
                    st.download_button("⬇️ Download Word (.docx)", data=docx_bytes,
                                       file_name="funding_report.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    html_bytes = generate_html_report(df_for_export, st.session_state.notes, title=report_title)
                    st.warning("`python-docx` not found. Generated HTML instead; open it and use your browser’s Print → Save as PDF.")
                    st.download_button("⬇️ Download HTML", data=html_bytes,
                                       file_name="funding_report.html",
                                       mime="text/html")
            except Exception as e:
                st.error(f"Failed to generate report: {e}")
    else:
        st.info("Select at least one call above to add notes and generate a report.")
