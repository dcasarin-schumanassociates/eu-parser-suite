# app_b3_5.py — Streamlit Funding Dashboard (Schuman-branded)
# Keeps app_b3_4 functionality; adds brand fonts/colors, hero header, and Altair theme.

from __future__ import annotations
import io, re, base64, os
from datetime import datetime
from typing import List, Dict
from pathlib import Path

import pandas as pd
import streamlit as st
import altair as alt

# Optional DOCX for shortlist export
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ------------------------------------------------------------
# Brand assets locations (put files here in your repo)
ASSETS_DIR = Path("assets")
FONTS_DIR  = ASSETS_DIR / "fonts"
LOGO_BLUE  = ASSETS_DIR / "logo-schuman_blue.png"           # primary
LOGO_GREY  = ASSETS_DIR / "logo-schuman_grey.png"           # optional
LOGO_WHITE = ASSETS_DIR / "logo-schuman_white.png"          # optional
LOGO_SMALL = ASSETS_DIR / "logo-schuman-blue-white-bg-SMALL.png"  # for page icon (optional)
# Font files you uploaded (place them in /assets/fonts/)
FONT_FILES = [
    ("SA Brand", "normal", 300, ASSETS_DIR / "Frm-Light.otf"),
    ("SA Brand", "normal", 400, ASSETS_DIR / "Frm-Regular.otf"),
    ("SA Brand", "normal", 500, ASSETS_DIR / "Frm-Medium.otf"),
    ("SA Brand", "normal", 700, ASSETS_DIR / "Frm-Bold.otf"),
    ("SA Brand", "italic", 700, ASSETS_DIR / "Frm-Bold-Italic.otf"),
    ("SA Brand", "normal", 900, ASSETS_DIR / "Frm-Black.otf"),
]

def _file_to_base64(p: Path) -> str | None:
    try:
        return base64.b64encode(p.read_bytes()).decode("utf-8")
    except Exception:
        return None

def inject_brand_css():
    """Inject global CSS with brand palette, fonts, rounded corners, shadows."""
    # Build @font-face blocks for any available OTFs
    font_faces = []
    for fam, style, weight, path in FONT_FILES:
        if path.exists():
            b64 = _file_to_base64(path)
            if b64:
                font_faces.append(f"""
                @font-face {{
                  font-family: '{fam}';
                  src: url(data:font/otf;base64,{b64}) format('opentype');
                  font-style: {style};
                  font-weight: {weight};
                  font-display: swap;
                }}
                """)
    font_css = "\n".join(font_faces)

    # Inject combined CSS
    st.markdown(f"""
    <style>
      {font_css}
      :root {{
        --sa-primary:   #1E4F86;
        --sa-primary-600:#17406B;
        --sa-primary-700:#123454;
        --sa-accent:    #00B4D8;
        --sa-ink:       #0F172A;
        --sa-muted:     #64748B;
        --sa-bg:        #FFFFFF;
        --sa-surface:   #F7F9FF;
        --sa-border:    #E5E7EB;
        --sa-font: 'SA Brand', system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, 'Apple Color Emoji','Segoe UI Emoji', sans-serif;
        --sa-radius: 12px;
        --sa-shadow: 0 6px 24px rgba(23,32,84,0.08);
      }}

      html, body, .stApp {{
        font-family: var(--sa-font) !important;
        color: var(--sa-ink);
        background: var(--sa-bg);
      }}

      .main .block-container {{ max-width: 1200px; }}
      h1,h2,h3,.stMarkdown h1,.stMarkdown h2,.stMarkdown h3{{ letter-spacing:.2px; font-weight:700; }}

      /* ---- FONT FAMILY + WEIGHTS FOR WIDGETS ---- */
      .stTextInput input,
      .stNumberInput input,
      .stTextArea textarea {{
        font-family: 'SA Brand', system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif !important;
        font-weight: 400;
      }}
      [data-testid="stSelectbox"] *,
      [data-testid="stMultiSelect"] *,
      div[role="listbox"] * {{
        font-family: 'SA Brand', system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif !important;
        font-weight: 500;
      }}
      [data-testid="stSlider"] *,
      .stSlider {{
        font-family: 'SA Brand', system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif !important;
        font-weight: 500;
      }}
      [data-testid="stRadio"] label,
      [data-testid="stCheckbox"] label {{
        font-family: 'SA Brand', system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif !important;
        font-weight: 500;
      }}
      .stForm h1, .stForm h2, .stForm h3, .stForm .stMarkdown h2 {{
        font-family: 'SA Brand', system-ui, -apple-system, Segoe UI, Roboto, Helvetica, Arial, sans-serif !important;
        font-weight: 700;
        letter-spacing: .2px;
      }}

      /* Buttons */
      .stButton button{{
        border-radius: var(--sa-radius);
        border: 1px solid var(--sa-primary-600);
        background: var(--sa-primary);
        color: #fff; font-weight: 600;
        box-shadow: var(--sa-shadow);
      }}
      .stButton button:hover{{ background: var(--sa-primary-600); border-color: var(--sa-primary-700); }}

      /* Tabs */
      [role="tablist"] {{ gap: 6px; }}
      [role="tab"]{{
        border: 1px solid var(--sa-border);
        background: var(--sa-surface);
        border-radius: var(--sa-radius);
        padding: .4rem .8rem;
        font-weight: 600;
      }}
      [role="tab"][aria-selected="true"]{{
        border-color: var(--sa-primary-600);
        box-shadow: var(--sa-shadow);
        background: #fff;
      }}

      /* Cards / Scroll containers */
      .scroll-container{{
        overflow:auto; max-height: 900px; padding:16px; border:1px solid var(--sa-border);
        border-radius: var(--sa-radius); background: #fff; box-shadow: var(--sa-shadow);
      }}

      /* Inputs - roundness */
      .stMultiSelect, .stSelectbox, .stTextInput, .stSlider {{ border-radius: var(--sa-radius) !important; }}
      .stDataFrame, .stTable {{ font-size: .94rem; }}
    </style>
    """, unsafe_allow_html=True)


def brand_header():
    """Top hero header with gradient and logo."""
    logo_src = None
    if LOGO_WHITE.exists():
        logo_src = f"data:image/png;base64,{_file_to_base64(LOGO_WHITE)}"
    elif LOGO_GREY.exists():
        logo_src = f"data:image/png;base64,{_file_to_base64(LOGO_GREY)}"
    elif LOGO_BLUE.exists():
        logo_src = f"data:image/png;base64,{_file_to_base64(LOGO_BLUE)}"

    st.markdown(f"""
    <div style="
      border-radius: 16px;
      background: linear-gradient(90deg, var(--sa-primary) 0%, var(--sa-primary-600) 65%, var(--sa-accent) 100%);
      padding: 18px 20px; color: white; display:flex; align-items:center; gap:16px;
      box-shadow: var(--sa-shadow);">
      {'<img src="'+logo_src+'" alt="Schuman Associates" style="height:44px; filter:brightness(1.05) contrast(1.05);" />' if logo_src else ''}
      <div style="flex:1;">
        <div style="font-size:18px; opacity:.95; font-weight:700;">Schuman Associates · Funding Dashboard</div>
        <div style="font-size:13px; opacity:.9;">Your European partners in a global market since 1989</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

# ------------------------------------------------------------
# Column mapping / helpers / canonicalisation (unchanged logic)
COLUMN_MAP = {
    "Programme": "programme",
    "Code": "code",
    "Title": "title",
    "Opening Date": "opening_date",
    "Opening date": "opening_date",
    "Deadline": "deadline",
    "First Stage Deadline": "first_deadline",
    "Second Stage Deadline": "second_deadline",
    "Second Stage deadline": "second_deadline",
    "Two-Stage": "two_stage",
    "Cluster": "cluster",
    "Destination": "destination",
    "Destination / Strand": "destination",
    "Destination/Strand": "destination",
    "Strand": "destination",
    "Budget Per Project": "budget_per_project_eur",
    "Budget per project": "budget_per_project_eur",
    "Total Budget": "total_budget_eur",
    "Number of Projects": "num_projects",
    "Type of Action": "type_of_action",
    "TRL": "trl",
    "Call Name": "call_name",
    "Expected Outcome": "expected_outcome",
    "Scope": "scope",
    "Description": "full_text",
    "Source Filename": "source_filename",
    "Version Label": "version_label",
    "Parsed On (UTC)": "parsed_on_utc",
    # Erasmus-specific:
    "Managing Authority": "managing_authority",
    "Key Action": "key_action",
}
DISPLAY_COLS = [
    "programme","code","title","opening_date","deadline",
    "type_of_action","budget_per_project_eur",
    "cluster","destination","trl",
    "managing_authority","key_action",
    "first_deadline","second_deadline","two_stage",
    "call_name","version_label","source_filename","parsed_on_utc",
]
SEARCHABLE_COLUMNS = (
    "code","title","call_name","expected_outcome","scope","full_text",
    "cluster","destination","type_of_action","trl","managing_authority","key_action"
)

def nl_to_br(s: str) -> str:
    return "" if not s else s.replace("\n", "<br>")

def clean_footer(text: str) -> str:
    if not text:
        return ""
    pat = re.compile(r"Horizon\s*Europe\s*[-–]?\s*Work Programme.*?Page\s+\d+\s+of\s+\d+", re.IGNORECASE | re.DOTALL)
    cleaned = pat.sub("", text)
    return re.sub(r"\s+", " ", cleaned).strip()

def normalize_bullets(text: str) -> str:
    if not isinstance(text, str) or text == "":
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(?m)^[ \t]*[▪◦●•]\s*", "- ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"(?<!\n)([ \t]+[-*]\s+)", r"\n- ", text)
    text = re.sub(r"(?<!\n)([ \t]+)(\d+\.\s+)", r"\n\2", text)
    return text.strip()

def highlight_text(text: str, keywords: list[str], colours=None) -> str:
    if not text:
        return ""
    kws = [str(k).strip() for k in keywords if k and str(k).strip()]
    if not kws:
        return text
    if colours is None:
        colours = ["#ffff00", "#a0e7e5", "#ffb3b3"]
    out = str(text)
    for i, kw in enumerate(kws):
        colour = colours[i % len(colours)]
        out = re.sub(re.escape(kw), lambda m: f"<span style='background-color:{colour}; font-weight:bold;'>{m.group(0)}</span>", out, flags=re.IGNORECASE)
    return out

def safe_date_series(s: pd.Series) -> pd.Series:
    out = pd.to_datetime(s, errors="coerce", dayfirst=True)
    if out.notna().sum() == 0:
        out = pd.to_datetime(s, errors="coerce", dayfirst=False)
    return out

def canonicalise(df: pd.DataFrame, programme_name: str) -> pd.DataFrame:
    df.columns = [c.strip() for c in df.columns]
    for src, dst in COLUMN_MAP.items():
        if src in df.columns and dst not in df.columns:
            df = df.rename(columns={src: dst})
    df = df.rename(columns={c: c.strip().lower() for c in df.columns})

    df["programme"] = programme_name

    for c in ("budget_per_project_eur","total_budget_eur","trl","num_projects"):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    for c in ("opening_date","deadline","first_deadline","second_deadline"):
        if c in df.columns:
            df[c] = safe_date_series(df[c])

    if "two_stage" in df.columns:
        df["two_stage"] = (
            df["two_stage"].astype(str).str.strip().str.lower()
            .map({"true": True, "false": False, "yes": True, "no": False, "1": True, "0": False})
            .fillna(False)
        )
    else:
        df["two_stage"] = False

    present = [c for c in SEARCHABLE_COLUMNS if c in df.columns]
    df["_search_all"]   = df[present].astype(str).agg(" ".join, axis=1).str.lower() if present else ""
    title_cols = [c for c in ["code","title"] if c in df.columns]
    df["_search_title"] = df[title_cols].astype(str).agg(" ".join, axis=1).str.lower() if title_cols else ""

    close_cols = [c for c in ["deadline","first_deadline","second_deadline"] if c in df.columns]
    if close_cols:
        df["closing_date_any"] = pd.to_datetime(df[close_cols].stack(), errors="coerce").groupby(level=0).min()
    else:
        df["closing_date_any"] = pd.NaT
    df["opening_year"]  = df["opening_date"].dt.year
    df["deadline_year"] = df["deadline"].dt.year

    return df

# --------- Caching ----------
@st.cache_data(show_spinner=False)
def get_sheet_names(file_bytes: bytes) -> List[str]:
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    return xls.sheet_names

@st.cache_data(show_spinner=False)
def load_programme(file_bytes: bytes, sheet_name: str, programme_name: str, _ver:int=1) -> pd.DataFrame:
    _ = hash(file_bytes)  # cache key includes content
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    raw = pd.read_excel(xls, sheet_name=sheet_name)
    df = canonicalise(raw, programme_name)
    return df.copy(deep=True)

# --------- Chart prep ----------
def wrap_label(text: str, width=60, max_lines=3) -> str:
    s = str(text or "")
    chunks = [s[i:i+width] for i in range(0, len(s), width)]
    return "\n".join(chunks[:max_lines]) if chunks else "—"

def build_month_bands(min_x: pd.Timestamp, max_x: pd.Timestamp) -> pd.DataFrame:
    start = pd.Timestamp(min_x).to_period("M").start_time
    end   = (pd.Timestamp(max_x).to_period("M") + 1).start_time
    months = pd.date_range(start, end, freq="MS")
    return pd.DataFrame({"start": months[:-1], "end": months[1:], "band": [i % 2 for i in range(len(months)-1)]})

def build_singlebar_rows(df: pd.DataFrame) -> pd.DataFrame:
    g = df.copy()
    if "code" in g.columns and g["code"].notna().any():
        base = g["code"].fillna("").astype(str)
    elif "title" in g.columns and g["title"].notna().any():
        base = g["title"].fillna("").astype(str)
    else:
        base = pd.Series([f"row-{i}" for i in range(len(g))], index=g.index)
    if base.duplicated(keep=False).any():
        base = base + g.groupby(base).cumcount().astype(str).radd("#")

    g["y_label"] = base.map(lambda s: wrap_label(s, width=100, max_lines=5))
    g["title_inbar"] = g.get("title","").astype(str).map(lambda s: wrap_label(s, width=100, max_lines=3))
    g = g[pd.notna(g["opening_date"]) & pd.notna(g["deadline"]) & (g["opening_date"] <= g["deadline"])].copy()
    if g.empty:
        return g
    g["bar_days"] = (g["deadline"] - g["opening_date"]).dt.days
    g["mid"] = g["opening_date"] + (g["deadline"] - g["opening_date"])/2
    return g.sort_values(["deadline","opening_date"])

def sa_altair_theme():
    # Align charts with UI fonts & colors
    return {
        "config": {
            "font": "SA Brand",
            "axis":   {"labelFont": "SA Brand", "titleFont": "SA Brand", "labelColor":"#0F172A"},
            "legend": {"labelFont": "SA Brand", "titleFont": "SA Brand"},
            "header": {"labelFont": "SA Brand", "titleFont": "SA Brand"},
            "title":  {"font": "SA Brand", "fontSize": 16, "fontWeight": 700, "color":"#0F172A"},
            "range": {
                "category": ["#1E4F86","#66C2A5","#FC8D62","#8DA0CB","#E78AC3","#A6D854","#FFD92F","#E5C494","#B3B3B3"]
            },
            "view": {"stroke": "transparent"}
        }
    }

alt.themes.register("sa_theme", sa_altair_theme)
alt.themes.enable("sa_theme")

def gantt_singlebar_chart(g: pd.DataFrame, color_field: str = "type_of_action", title: str = ""):
    if g is None or g.empty:
        return None

    min_x = min(g["opening_date"].min(), g["deadline"].min())
    max_x = max(g["opening_date"].max(), g["deadline"].max())
    bands_df = build_month_bands(min_x, max_x)

    month_shade = alt.Chart(bands_df).mark_rect(tooltip=False).encode(
        x="start:T", x2="end:T",
        opacity=alt.Opacity("band:Q", scale=alt.Scale(domain=[0,1], range=[0.0,0.05]), legend=None),
        color=alt.value("#1E4F86")
    )
    months = pd.date_range(pd.Timestamp(min_x).to_period("M").start_time,
                           pd.Timestamp(max_x).to_period("M").end_time, freq="MS")
    month_grid = alt.Chart(pd.DataFrame({"t": months})).mark_rule(stroke="#1E4F86", strokeWidth=0.3).encode(x="t:T")
    month_labels_df = pd.DataFrame({
        "month": months[:-1], "next_month": months[1:],
        "label": [m.strftime("%b %Y") for m in months[:-1]]
    })
    month_labels_df["mid"] = month_labels_df["month"] + ((month_labels_df["next_month"] - month_labels_df["month"]) / 2)
    month_labels = alt.Chart(month_labels_df).mark_text(
        align="center", baseline="top", dy=0, fontSize=12, fontWeight="bold"
    ).encode(x="mid:T", text="label:N", y=alt.value(0))

    today_ts = pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None)
    today_df = pd.DataFrame({"t":[today_ts]})
    today_rule = alt.Chart(today_df).mark_rule(color="#1E4F86", strokeDash=[2,1], strokeWidth=2).encode(
        x="t:T", tooltip=[alt.Tooltip("t:T", title="Today", format="%d %b %Y")]
    )
    today_label = alt.Chart(today_df).mark_text(
        align="left", baseline="top", dx=4, dy=18, fontSize=11, fontWeight="bold", color="#1E4F86"
    ).encode(x="t:T", y=alt.value(0), text=alt.Text("t:T", format='Today: "%d %b %Y"'))

    y_order = g["y_label"].drop_duplicates().tolist()
    row_h = 46
    bar_size = int(row_h * 0.38)
    domain_min, domain_max = g["opening_date"].min(), g["deadline"].max()

    base = alt.Chart(g).encode(
        y=alt.Y("y_label:N", sort=y_order,
                axis=alt.Axis(title=None, labelLimit=200, labelFontSize=11, labelAlign="right", labelPadding=50, domain=True),
                scale=alt.Scale(domain=y_order, paddingInner=0.3, paddingOuter=0.8))
    )

    bars = base.mark_bar(cornerRadius=10, size=bar_size).encode(
        x=alt.X("opening_date:T",
                axis=alt.Axis(title=None, format="%b %Y", tickCount="month", orient="top",
                              labelFontSize=11, labelPadding=50, labelOverlap="greedy", tickSize=6),
                scale=alt.Scale(domain=[domain_min, domain_max])),
        x2="deadline:T",
        color=alt.Color(color_field + ":N",
                        legend=alt.Legend(title=color_field.replace("_"," ").title(),
                                          orient="top", direction="horizontal", offset=100),
                        scale=alt.Scale(scheme="set2")),
        tooltip=[
            alt.Tooltip("title:N", title="Title"),
            alt.Tooltip("opening_date:T", title="Opening", format="%d %b %Y"),
            alt.Tooltip("deadline:T", title="Deadline", format="%d %b %Y"),
            alt.Tooltip(color_field + ":N", title=color_field.replace("_"," ").title()),
        ]
    )

    start_labels = base.mark_text(align="right", dx=-4, dy=5, fontSize=10, color="#111")\
        .encode(x="opening_date:T", text=alt.Text("opening_date:T", format="%d %b %Y"))
    end_labels   = base.mark_text(align="left",  dx=4, dy=5, fontSize=10, color="#111")\
        .encode(x="deadline:T",      text=alt.Text("deadline:T",      format="%d %b %Y"))
    inbar = base.mark_text(align="left", baseline="bottom", dx=2, dy=-(int(bar_size/2)+4), color="black")\
        .encode(x=alt.X("opening_date:T", scale=alt.Scale(domain=[domain_min, domain_max]), axis=None),
                text="title_inbar:N",
                opacity=alt.condition(alt.datum.bar_days >= 10, alt.value(1), alt.value(0)))

    chart = (month_shade + month_grid + bars + start_labels + end_labels + inbar + month_labels + today_rule + today_label)\
        .properties(height=max(800, len(y_order)*row_h), width='container',
                    padding={"top":50,"bottom":30,"left":10,"right":10})\
        .configure_axis(grid=False, domain=True, domainWidth=1)\
        .configure_view(continuousHeight=500, continuousWidth=500, strokeWidth=0, clip=False)\
        .interactive(bind_x=True)

    return chart if not title else chart.properties(title=title)

# --------- Report (DOCX) ----------
def generate_docx_report(calls_df: pd.DataFrame, notes_by_code: Dict[str,str], title="Funding Report") -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    h = doc.add_heading(title, level=0); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(f"Generated on {datetime.utcnow():%d %b %Y, %H:%M UTC}"); p.runs[0].font.size = Pt(9)

    table = doc.add_table(rows=1, cols=5); hdr = table.rows[0].cells
    for i, t in enumerate(["Programme","Code","Title","Opening","Deadline"]): hdr[i].text = t
    for _, r in calls_df.iterrows():
        row = table.add_row().cells
        row[0].text = str(r.get("programme","")); row[1].text = str(r.get("code","")); row[2].text = str(r.get("title",""))
        op, dl = r.get("opening_date"), r.get("deadline")
        row[3].text = op.strftime("%d %b %Y") if pd.notna(op) else "-"
        row[4].text = dl.strftime("%d %b %Y") if pd.notna(dl) else "-"

    for _, r in calls_df.iterrows():
        doc.add_page_break()
        doc.add_heading(f"{r.get('code','')} — {r.get('title','')}", level=1)
        lines = []
        lines.append(f"Programme: {r.get('programme','-')}")
        lines.append(f"Cluster: {r.get('cluster','-')}")
        lines.append(f"Destination: {r.get('destination','-')}")
        lines.append(f"Type of Action: {r.get('type_of_action','-')}")
        trl_val = r.get("trl"); lines.append(f"TRL: {int(trl_val) if pd.notna(trl_val) else '-'}")
        ma = r.get("managing_authority"); ka = r.get("key_action")
        if pd.notna(ma) or pd.notna(ka):
            lines.append(f"Managing Authority: {ma if pd.notna(ma) else '-'}")
            lines.append(f"Key Action: {ka if pd.notna(ka) else '-'}")
        op, dl = r.get("opening_date"), r.get("deadline")
        lines.append(f"Opening: {op:%d %b %Y}" if pd.notna(op) else "Opening: -")
        lines.append(f"Deadline: {dl:%d %b %Y}" if pd.notna(dl) else "Deadline: -")
        doc.add_paragraph("\n".join(lines))
        notes = (notes_by_code or {}).get(str(r.get("code","")), "")
        doc.add_heading("Notes", level=2); doc.add_paragraph(notes if notes else "-")

    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.getvalue()

# ------------------------------------------------------------
# UI — page config + branding
page_icon = str(LOGO_SMALL) if LOGO_SMALL.exists() else "🟦"
st.set_page_config(page_title="Schuman · Funding Dashboard", page_icon=page_icon, layout="wide")
inject_brand_css()
brand_header()

st.info(
    "📂 Please upload the latest parsed Excel file.\n\n"
    "➡️ Location hint:\n\n"
    "- **3.SA Practices** → Central Systems and Bid Management → 1. Central Systems\n\n"
    "👉 Look for *Central System Funding Compass Database*.\n"
)

upl = st.file_uploader("Upload Excel (.xlsx)", type=["xlsx"])
if not upl:
    st.stop()

# Detect sheets and allow override
sheets = get_sheet_names(upl.getvalue())
c1, c2 = st.columns(2)
with c1:
    hz_sheet = st.selectbox("Horizon Database", options=sheets, index=0)
with c2:
    er_sheet = st.selectbox("Erasmus Database", options=sheets, index=min(1, len(sheets)-1))

# Load each programme independently
df_h = load_programme(upl.getvalue(), hz_sheet, "Horizon Europe", _ver=1)
df_e = load_programme(upl.getvalue(), er_sheet, "Erasmus+",      _ver=1)

# ------- Build filter choices -------
all_df = pd.concat([df_h.assign(programme="Horizon Europe"), df_e.assign(programme="Erasmus+")], ignore_index=True)
opening_years  = sorted([int(y) for y in all_df["opening_year"].dropna().unique()])
deadline_years = sorted([int(y) for y in all_df["deadline_year"].dropna().unique()])
type_opts      = sorted([t for t in all_df.get("type_of_action", pd.Series(dtype=object)).dropna().unique().tolist() if t!=""])

cluster_opts   = sorted([c for c in df_h.get("cluster", pd.Series(dtype=object)).dropna().unique().tolist() if c!=""])
dest_opts      = sorted([d for d in df_h.get("destination", pd.Series(dtype=object)).dropna().unique().tolist() if d!=""])
trl_opts       = sorted([str(int(x)) for x in df_h.get("trl", pd.Series(dtype=float)).dropna().unique() if pd.notna(x)])

ma_opts        = sorted([m for m in df_e.get("managing_authority", pd.Series(dtype=object)).dropna().unique().tolist() if m!=""])
ka_opts        = sorted([k for k in df_e.get("key_action", pd.Series(dtype=object)).dropna().unique().tolist() if k!=""])

# Budget slider (use combined range)
bud_series = pd.to_numeric(all_df.get("budget_per_project_eur"), errors="coerce").dropna()
if bud_series.empty:
    min_bud, max_bud = 0.0, 1_000_000.0
else:
    min_bud, max_bud = float(bud_series.min()), float(bud_series.max())
    if not (min_bud < max_bud):
        min_bud, max_bud = max(min_bud, 0.0), min_bud + 100000.0
rng = max_bud - min_bud
step = max(1e4, round(rng / 50, -3)) if rng else 10000.0

with st.form("filters", clear_on_submit=False):
    st.header("Filters")

    # Row 1: opening, deadline, open calls
    col_oy, col_dy, col_open = st.columns([1,1,1])
    with col_oy:
        open_year_sel = st.multiselect("Opening year(s)", opening_years, default=opening_years)
    with col_dy:
        deadline_year_sel = st.multiselect("Deadline year(s)", deadline_years, default=deadline_years)
    with col_open:
        open_calls_only = st.checkbox(
            "Open calls only",
            value=False,
            help="Keep only entries whose final Deadline is strictly after today (Europe/Brussels)."
        )

    # Row 2: budget
    col_bud, _sp = st.columns([2,1])
    with col_bud:
        budget_range = st.slider("Budget per project (EUR)", min_bud, max_bud, (min_bud, max_bud), step=step)

    # Row 3: keywords + combine + title/code toggle
    k1, k2, k3, kcomb, ktit = st.columns([2,2,2,1,1.2])
    with k1: kw1 = st.text_input("Keyword 1")
    with k2: kw2 = st.text_input("Keyword 2")
    with k3: kw3 = st.text_input("Keyword 3")
    with kcomb:
        kw_mode = st.radio("Combine", ["AND","OR"], index=1, horizontal=True)  # default OR
    with ktit:
        title_code_only = st.checkbox("Title/Code only", value=False)          # default off

    # Row 4: type of action
    types_sel = st.multiselect("Type of Action", type_opts)

    # Horizon-specific
    st.subheader("Horizon-specific")
    h1,h2,h3 = st.columns(3)
    with h1: clusters_sel = st.multiselect("Cluster", cluster_opts)
    with h2: dests_sel    = st.multiselect("Destination", dest_opts)
    with h3: trls_sel     = st.multiselect("TRL", trl_opts)

    # Erasmus-specific
    st.subheader("Erasmus+-specific")
    e1,e2 = st.columns(2)
    with e1: ma_sel = st.multiselect("Managing Authority", ma_opts)
    with e2: ka_sel = st.multiselect("Key Action", ka_opts)

    applied = st.form_submit_button("Apply filters")

# Persist criteria
if "crit35" not in st.session_state:
    st.session_state.crit35 = {}
if applied or not st.session_state.crit35:
    st.session_state.crit35 = dict(
        open_years=open_year_sel, deadline_years=deadline_year_sel,
        types=types_sel, kws=[kw1,kw2,kw3], kw_mode=kw_mode, title_code_only=title_code_only,
        budget_range=budget_range,
        clusters=clusters_sel, dests=dests_sel, trls=trls_sel,
        managing_authority=ma_sel, key_action=ka_sel,
        open_calls_only=open_calls_only
    )
crit = st.session_state.crit35

# Filtering helpers (unchanged)
def multi_keyword_filter(df: pd.DataFrame, terms: list[str], mode: str, title_code_only: bool) -> pd.DataFrame:
    terms = [t.strip().lower() for t in terms if t and str(t).strip()]
    if not terms:
        return df
    hay = df["_search_title"] if title_code_only else df["_search_all"]
    if mode.upper() == "AND":
        pattern = "".join([f"(?=.*{re.escape(t)})" for t in terms]) + ".*"
    else:
        pattern = "(" + "|".join(re.escape(t) for t in terms) + ")"
    return df[hay.str_contains(pattern, regex=True, na=False)] if hasattr(pd.Series, "str_contains") else df[hay.str.contains(pattern, regex=True, na=False)]

def apply_common_filters(df0: pd.DataFrame) -> pd.DataFrame:
    df = df0.copy()
    if crit["open_years"]:
        df = df[df["opening_year"].isin(crit["open_years"])]
    if crit["deadline_years"]:
        df = df[df["deadline_year"].isin(crit["deadline_years"])]
    if crit.get("open_calls_only"):
        today = pd.Timestamp.now(tz="Europe/Brussels").normalize().tz_localize(None)
        df = df[df["deadline"].notna() & (df["deadline"] > today)]
    if crit["types"]:
        df = df[df.get("type_of_action").isin(crit["types"])]
    lo, hi = crit["budget_range"]
    df = df[df.get("budget_per_project_eur").fillna(0).between(lo, hi)]
    df = multi_keyword_filter(df, crit["kws"], crit["kw_mode"], crit["title_code_only"])
    return df

def apply_horizon_filters(df0: pd.DataFrame) -> pd.DataFrame:
    df = apply_common_filters(df0)
    if crit["clusters"]: df = df[df.get("cluster").isin(crit["clusters"])]
    if crit["dests"]:    df = df[df.get("destination").isin(crit["dests"])]
    if crit["trls"]:
        df = df[df.get("trl").dropna().astype("Int64").astype(str).isin(crit["trls"])]
    return df

def apply_erasmus_filters(df0: pd.DataFrame) -> pd.DataFrame:
    df = apply_common_filters(df0)
    if crit["managing_authority"]: df = df[df.get("managing_authority").isin(crit["managing_authority"])]
    if crit["key_action"]:         df = df[df.get("key_action").isin(crit["key_action"])]
    return df

fh = apply_horizon_filters(df_h)
fe = apply_erasmus_filters(df_e)
st.caption(f"Rows after filters — Horizon: {len(fh)} | Erasmus: {len(fe)}")

# ------------------------------ Tabs ------------------------------
tab_hz, tab_er, tab_tbl, tab_full, tab_short = st.tabs([
    "📅 Gantt — Horizon", "📅 Gantt — Erasmus", "📋 Tables", "📚 Full Data", "📝 Shortlist"
])

with tab_hz:
    st.subheader("Gantt — Horizon Europe (Opening → Deadline)")
    group_by_cluster = st.checkbox(
        "Group by cluster (render one Gantt per cluster)",
        value=False,
        help="When enabled, the Horizon chart is split into one dropdown per Cluster."
    )
    from copy import deepcopy
    if not group_by_cluster:
        g_h = build_singlebar_rows(fh)
        if g_h.empty:
            st.info("No valid Horizon rows/dates.")
        else:
            st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
            st.altair_chart(gantt_singlebar_chart(g_h, color_field="type_of_action"), use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)
    else:
        if "cluster" not in fh.columns:
            st.warning("Column 'cluster' not available in Horizon dataset.")
        else:
            tmp = fh.copy()
            tmp["cluster"] = tmp["cluster"].fillna("— Unspecified —").replace({"": "— Unspecified —"})
            groups = list(tmp.groupby("cluster", dropna=False))
            groups.sort(key=lambda kv: len(kv[1]), reverse=True)
            st.caption(f"Clusters found: {len(groups)}")
            for clu_name, gdf in groups:
                g_clu = build_singlebar_rows(gdf)
                with st.expander(f"Cluster: {clu_name}  ({len(g_clu)} calls)", expanded=False):
                    if g_clu.empty:
                        st.info("No valid rows/dates for this cluster after filters.")
                    else:
                        st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
                        st.altair_chart(gantt_singlebar_chart(g_clu, color_field="type_of_action"), use_container_width=True)
                        st.markdown('</div>', unsafe_allow_html=True)

with tab_er:
    st.subheader("Gantt — Erasmus+ (Opening → Deadline)")
    g_e = build_singlebar_rows(fe)
    if g_e.empty: st.info("No valid Erasmus rows/dates.")
    else:
        st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
        st.altair_chart(gantt_singlebar_chart(g_e, color_field="type_of_action"), use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

with tab_tbl:
    st.subheader("Tables")
    show_cols_h = [c for c in DISPLAY_COLS if c in fh.columns]
    show_cols_e = [c for c in DISPLAY_COLS if c in fe.columns]
    colA, colB = st.columns(2)
    with colA:
        st.markdown(f"### Horizon Europe ({len(fh)})")
        if len(fh): st.dataframe(fh[show_cols_h], use_container_width=True, hide_index=True)
        else: st.caption("— no rows —")
    with colB:
        st.markdown(f"### Erasmus+ ({len(fe)})")
        if len(fe): st.dataframe(fe[show_cols_e], use_container_width=True, hide_index=True)
        else: st.caption("— no rows —")

with tab_full:
    st.subheader("Full Data — Expand rows for details")

    kw_list = crit.get("kws", [])

    def render_row(row, programme: str):
        c1, c2 = st.columns(2)
        with c1:
            if pd.notna(row.get("opening_date")):
                st.markdown(f"📅 **Opening:** {row['opening_date']:%d %b %Y}")
            if pd.notna(row.get("deadline")):
                st.markdown(f"⏳ **Deadline:** {row['deadline']:%d %b %Y}")
        with c2:
            if pd.notna(row.get("budget_per_project_eur")):
                st.markdown(f"💶 **Budget/Project:** {row['budget_per_project_eur']:,.0f} EUR")
            if pd.notna(row.get("total_budget_eur")):
                st.markdown(f"📦 **Total:** {row['total_budget_eur']:,.0f} EUR")
            if pd.notna(row.get("num_projects")):
                st.markdown(f"📊 **# Projects:** {int(row['num_projects'])}")

        meta_bits = [
            f"🏷️ **Programme:** {programme}",
            f"**Type of Action:** {row.get('type_of_action','-')}",
        ]
        if programme == "Horizon Europe":
            meta_bits += [
                f"**Cluster:** {row.get('cluster','-')}",
                f"**Destination:** {row.get('destination','-')}",
                f"**TRL:** {row.get('trl','-')}",
            ]
        else:
            meta_bits += [
                f"**Managing Authority:** {row.get('managing_authority','-')}",
                f"**Key Action:** {row.get('key_action','-')}",
            ]
        st.markdown(" | ".join(meta_bits))

        if row.get("expected_outcome"):
            with st.expander("🎯 Expected Outcome"):
                clean_text = nl_to_br(normalize_bullets(clean_footer(str(row.get("expected_outcome")))))
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)
        if row.get("scope"):
            with st.expander("🧭 Scope"):
                clean_text = nl_to_br(normalize_bullets(clean_footer(str(row.get("scope")))))
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)
        if row.get("full_text"):
            with st.expander("📖 Full Description"):
                clean_text = nl_to_br(normalize_bullets(clean_footer(str(row.get("full_text")))))
                st.markdown(highlight_text(clean_text, kw_list), unsafe_allow_html=True)

        st.caption(
            f"📂 Source: {row.get('source_filename','-')} "
            f"| Version: {row.get('version_label','-')} "
            f"| Parsed on: {row.get('parsed_on_utc','-')}"
        )

    st.markdown(f"### Horizon Europe ({len(fh)})")
    group_hz = st.checkbox(
        "Group Horizon by Cluster (dropdowns)",
        value=False,
        help="Show one expander per Cluster; inside each, expand rows for details."
    )
    if len(fh) == 0:
        st.caption("— no Horizon rows after filters —")
    else:
        if not group_hz:
            for _, r in fh.iterrows():
                label = f"{str(r.get('code') or '')} — {str(r.get('title') or '')}".strip(" —")
                with st.expander(label or "(untitled)"):
                    render_row(r, "Horizon Europe")
        else:
            tmp = fh.copy()
            tmp["cluster"] = tmp.get("cluster").fillna("— Unspecified —").replace({"": "— Unspecified —"})
            groups = list(tmp.groupby("cluster", dropna=False))
            groups.sort(key=lambda kv: len(kv[1]), reverse=True)
            st.caption(f"Clusters found: {len(groups)}")
            for clu_name, gdf in groups:
                disp = str(clu_name)
                with st.expander(f"Cluster: {disp}  ({len(gdf)} calls)", expanded=False):
                    for _, r in gdf.iterrows():
                        label = f"{str(r.get('code') or '')} — {str(r.get('title') or '')}".strip(" —")
                        with st.expander(label or "(untitled)"):
                            render_row(r, "Horizon Europe")

    st.markdown(f"### Erasmus+ ({len(fe)})")
    if len(fe) == 0:
        st.caption("— no Erasmus rows after filters —")
    else:
        for _, r in fe.iterrows():
            label = f"{str(r.get('code') or '')} — {str(r.get('title') or '')}".strip(" —")
            with st.expander(label or "(untitled)"):
                render_row(r, "Erasmus+")

with tab_short:
    st.subheader("Shortlist & Notes (DOCX)")
    if "sel35" not in st.session_state: st.session_state.sel35 = set()
    if "notes35" not in st.session_state: st.session_state.notes35 = {}

    combined = []
    if len(fh): combined.append(fh.assign(programme="Horizon Europe"))
    if len(fe): combined.append(fe.assign(programme="Erasmus+"))
    ff = pd.concat(combined, ignore_index=True) if combined else pd.DataFrame()

    st.markdown("**Select calls**")
    if "closing_date_any" not in ff.columns:
        close_cols = [c for c in ["deadline","first_deadline","second_deadline"] if c in ff.columns]
        if close_cols:
            ff["closing_date_any"] = pd.to_datetime(ff[close_cols].stack(), errors="coerce").groupby(level=0).min()
        else:
            ff["closing_date_any"] = pd.NaT
    sort_keys = [c for c in ["closing_date_any","opening_date"] if c in ff.columns]
    if sort_keys:
        ff = ff.sort_values(sort_keys)

    for idx, r in ff.iterrows():
        code = str(r.get("code") or ""); title = str(r.get("title") or "")
        label = f"{code} — {title}".strip(" —")
        checked = code in st.session_state.sel35
        new = st.checkbox(label or "(untitled)", value=checked, key=f"sel35_{code}_{idx}")
        if new and not checked: st.session_state.sel35.add(code)
        elif (not new) and checked: st.session_state.sel35.discard(code)

    selected_df = ff[ff["code"].astype(str).isin(st.session_state.sel35)]
    if not selected_df.empty:
        st.markdown("---")
        for _, r in selected_df.iterrows():
            code = str(r.get("code") or "")
            default = st.session_state.notes35.get(code, "")
            st.session_state.notes35[code] = st.text_area(f"Notes — {code}", value=default, height=110, key=f"note35_{code}")

        colA, colB = st.columns(2)
        with colA: title = st.text_input("Report title", value="Funding Report – Shortlist (app_b3.5)")
        with colB: pass

        if st.button("📄 Generate DOCX"):
            try:
                if DOCX_AVAILABLE:
                    data = generate_docx_report(selected_df, st.session_state.notes35, title=title)
                    st.download_button("⬇️ Download .docx", data=data,
                                       file_name="funding_report.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                else:
                    st.error("python-docx not installed in this environment.")
            except Exception as e:
                st.error(f"Failed to generate report: {e}")
    else:
        st.info("Select at least one call to add notes and export.")
